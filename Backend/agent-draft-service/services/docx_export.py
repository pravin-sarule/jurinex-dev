"""
Shared DOCX export: same format as the static preview (AssembledPreviewPage).
Used for: (1) Google Docs upload in assembler, (2) DOCX download.
Do not change static preview CSS; this module matches its layout so that
Google Docs and DOCX download look like the in-app static preview.

Static preview uses: width 210mm, min-height 297mm, padding 2.54cm, Times New Roman.
Here we use: A4 (21×29.7 cm), 2.54 cm margins, Times New Roman 12pt.
When template_css is provided, it is inlined into the HTML so DOCX/Google Docs
preserve the same formatting (alignment, indents, font) as the static preview.
When template_url is provided, styles are fetched from the template document
and merged with template_css so the Word/Google Docs output matches the
uploaded template format exactly.
"""

from __future__ import annotations

import io
import logging
import re
import zipfile
from xml.etree import ElementTree as ET
from bs4 import BeautifulSoup, Comment

logger = logging.getLogger(__name__)

# Match static preview .assembled-doc-container: A4, 2.54 cm, Times New Roman
FONT_NAME = "Times New Roman"
FONT_SIZE_PT = 12
PAGE_HEIGHT_CM = 29.7   # 297mm A4
PAGE_WIDTH_CM = 21     # 210mm A4
MARGIN_CM = 2.54
SECTION_BREAK_MARKER = "<!-- SECTION_BREAK -->"
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
PKG_REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
CONTENT_TYPES_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
ET.register_namespace("w", W_NS)
ET.register_namespace("r", R_NS)


def _extract_style_blocks(html: str) -> str:
    """Extract all <style>...</style> content from template HTML for format reference."""
    if not (html or "").strip():
        return ""
    # Match <style ...>...</style> (any attributes, multiline)
    blocks = re.findall(r"<style[^>]*>([\s\S]*?)</style>", html, re.IGNORECASE)
    return "\n".join(b.strip() for b in blocks if b.strip()).strip()


def _get_merged_template_css(template_css: str | None, template_url: str | None) -> str | None:
    """
    Merge template_css with styles fetched from template_url so conversion uses
    the exact format of the uploaded template (alignment, indents, fonts).
    """
    from services.template_format import fetch_template_html

    css_parts = []
    if (template_css or "").strip():
        css_parts.append(template_css.strip())
    if (template_url or "").strip():
        try:
            template_html = fetch_template_html(template_url.strip())
            if template_html:
                from_template = _extract_style_blocks(template_html)
                if from_template:
                    css_parts.append("/* styles from template document */\n" + from_template)
                    logger.info("DOCX export: merged template URL styles into CSS for format preservation")
        except Exception as e:
            logger.warning("DOCX export: could not fetch template URL for styles: %s", e)
    if not css_parts:
        return None
    return "\n\n".join(css_parts)


def _inline_template_css(html_content: str, template_css: str) -> str:
    """
    Inline template_css into the HTML so conversion to DOCX preserves the same
    format as the static preview (alignment, indents, font-size, etc.).
    Returns the body content with inlined styles (no <style> block in output).
    """
    if not (template_css or "").strip():
        return html_content
    try:
        from premailer import transform
    except ImportError:
        return html_content
    full_doc = (
        "<!DOCTYPE html><html><head><style>\n"
        + (template_css or "").strip()
        + "\n</style></head><body>\n"
        + html_content
        + "\n</body></html>"
    )
    try:
        inlined = transform(full_doc)
        # Extract body inner HTML (premailer returns full document)
        start = inlined.find("<body>")
        end = inlined.rfind("</body>")
        if start != -1 and end != -1:
            start += len("<body>")
            return inlined[start:end].strip()
    except Exception as e:
        import logging
        logging.getLogger(__name__).warning("premailer inline failed, using original HTML: %s", e)
    return html_content


def _strip_style_tags(html: str) -> str:
    """Remove <style>...</style> blocks so CSS is not shown as text in DOCX/Google Docs."""
    out = re.sub(r"<style[^>]*>[\s\S]*?</style>", "", html, flags=re.IGNORECASE)
    # Remove any stray CSS-like lines that might remain (e.g. from malformed tags)
    out = re.sub(r"@page\s*\{[^}]*\}", "", out, flags=re.IGNORECASE)
    out = re.sub(r"@media\s+print\s*\{[\s\S]*?\}", "", out, flags=re.IGNORECASE)
    out = re.sub(r"\.document-section\s*\{[^}]*\}", "", out, flags=re.IGNORECASE)
    out = re.sub(r"\.page-break\s*\{[^}]*\}", "", out, flags=re.IGNORECASE)
    return out.strip()


def _ensure_anchor_href(html: str) -> str:
    """Ensure every <a> tag has an href so htmldocx does not raise KeyError('href')."""
    def repl(m: re.Match) -> str:
        prefix, rest = m.group(1), m.group(2)
        if rest and "href=" in rest.lower():
            return m.group(0)
        if rest:
            return f'<a{prefix}href="#" {rest}>'
        return f'<a{prefix}href="#">'
    return re.sub(r"<a(\s*)([^>]*)>", repl, html, flags=re.IGNORECASE)


def _sanitize_html_for_docx(html: str) -> str:
    """
    Normalize fragile HTML before sending it to htmldocx.
    Keeps document structure, but removes wrapper/script/style noise and simplifies tables.
    """
    if not (html or "").strip():
        return ""

    soup = BeautifulSoup(html, "html.parser")

    for comment in soup.find_all(string=lambda value: isinstance(value, Comment)):
        comment.extract()

    for tag in soup.find_all(["script", "style", "meta", "link", "noscript"]):
        tag.decompose()

    for tag_name in ["html", "body", "head", "colgroup", "col", "thead", "tbody", "tfoot"]:
        for tag in soup.find_all(tag_name):
            tag.unwrap()

    for tag in soup.find_all(True):
        attrs = dict(tag.attrs or {})
        cleaned_attrs = {}
        for key, value in attrs.items():
            key_l = str(key).lower()
            if key_l.startswith("on") or key_l.startswith("data-"):
                continue
            if key_l in {"contenteditable", "spellcheck", "tabindex"}:
                continue
            cleaned_attrs[key] = value
        tag.attrs = cleaned_attrs

        if tag.name == "table":
            existing_style = tag.get("style", "")
            table_style = "border-collapse: collapse; width: 100%;"
            tag["style"] = f"{existing_style}; {table_style}".strip("; ").strip()

    return str(soup).strip()


def _plain_text_html(html: str) -> str:
    """Last-resort fallback so users still get a DOCX even if rich HTML conversion fails."""
    soup = BeautifulSoup(html or "", "html.parser")
    blocks = []
    for node in soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "p", "li", "td", "th"]):
        text = node.get_text(" ", strip=True)
        if text:
            blocks.append(f"<p>{text}</p>")
    if not blocks:
        text = soup.get_text("\n", strip=True)
        blocks = [f"<p>{line.strip()}</p>" for line in text.splitlines() if line.strip()]
    return "\n".join(blocks)


def _build_altchunk_docx(html: str) -> bytes:
    """Build a DOCX that embeds HTML directly, which preserves layout better in Word."""
    from docx import Document
    from docx.shared import Cm

    document = Document()
    try:
        section = document.sections[0]
        section.page_height = Cm(PAGE_HEIGHT_CM)
        section.page_width = Cm(PAGE_WIDTH_CM)
        section.left_margin = Cm(MARGIN_CM)
        section.right_margin = Cm(MARGIN_CM)
        section.top_margin = Cm(MARGIN_CM)
        section.bottom_margin = Cm(MARGIN_CM)
    except (IndexError, Exception) as _sec_err:
        logger.warning("_build_altchunk_docx: could not set page dimensions (%s), proceeding anyway", _sec_err)

    base_buf = io.BytesIO()
    document.save(base_buf)
    base_buf.seek(0)

    in_mem = io.BytesIO()
    with zipfile.ZipFile(base_buf, "r") as zin, zipfile.ZipFile(in_mem, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "word/document.xml":
                root = ET.fromstring(data)
                body = root.find(f"{{{W_NS}}}body")
                if body is None:
                    raise RuntimeError("DOCX body not found for altChunk export")
                sect_pr = body.find(f"{{{W_NS}}}sectPr")
                alt_chunk = ET.Element(f"{{{W_NS}}}altChunk", {f"{{{R_NS}}}id": "rIdAltChunkHtml"})
                if sect_pr is not None:
                    body.remove(sect_pr)
                    body.append(alt_chunk)
                    body.append(sect_pr)
                else:
                    body.append(alt_chunk)
                data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
            elif item.filename == "word/_rels/document.xml.rels":
                root = ET.fromstring(data)
                ET.SubElement(
                    root,
                    f"{{{PKG_REL_NS}}}Relationship",
                    {
                        "Id": "rIdAltChunkHtml",
                        "Type": "http://schemas.openxmlformats.org/officeDocument/2006/relationships/aFChunk",
                        "Target": "afchunk.html",
                    },
                )
                data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
            elif item.filename == "[Content_Types].xml":
                root = ET.fromstring(data)
                ET.SubElement(
                    root,
                    f"{{{CONTENT_TYPES_NS}}}Override",
                    {
                        "PartName": "/word/afchunk.html",
                        "ContentType": "text/html",
                    },
                )
                data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
            zout.writestr(item, data)

        zout.writestr("word/afchunk.html", html.encode("utf-8"))

    in_mem.seek(0)
    return in_mem.read()


def _normalize_google_import_html(html: str) -> str:
    """Prepare assembled HTML for direct Google Docs HTML import."""
    normalized = (html or "").replace(SECTION_BREAK_MARKER, '<div style="page-break-before: always;"></div>')
    normalized = re.sub(
        r"<table([^>]*)>",
        lambda match: (
            f'<table{match.group(1)} style="border-collapse: collapse; width: 100%; table-layout: fixed;">'
            if 'style=' not in match.group(1).lower()
            else f'<table{match.group(1)}>'
        ),
        normalized,
        flags=re.IGNORECASE,
    )
    normalized = re.sub(
        r"<(td|th)([^>]*)>",
        lambda match: (
            f'<{match.group(1)}{match.group(2)} style="border: 1px solid #000; vertical-align: top; padding: 6px;">'
            if 'style=' not in match.group(2).lower()
            else f'<{match.group(1)}{match.group(2)}>'
        ),
        normalized,
        flags=re.IGNORECASE,
    )
    return normalized


def assembled_html_to_google_import_html(
    html_content: str,
    template_css: str | None = None,
    template_url: str | None = None,
) -> str:
    """
    Build a single self-contained HTML document for direct Google Docs import.
    This avoids the extra DOCX->Google conversion step that can distort layout.
    """
    from services.assembled_doc_clean import strip_citations_for_assembled

    html_content = strip_citations_for_assembled(html_content or "")
    merged_css = _get_merged_template_css(template_css, template_url)
    if merged_css:
        html_content = _inline_template_css(html_content, merged_css)
    html_content = _ensure_anchor_href(_strip_style_tags(html_content))
    html_content = _normalize_google_import_html(html_content)

    base_css = """
        body {
            font-family: 'Times New Roman', Times, serif;
            font-size: 12pt;
            line-height: 1.5;
            color: #000;
            margin: 0;
        }
        p, li { margin: 0 0 0.5em 0; }
        table { border-collapse: collapse; width: 100%; table-layout: fixed; margin: 0.5em 0 1em; }
        td, th { border: 1px solid #000; vertical-align: top; padding: 6px; }
    """

    return (
        "<!DOCTYPE html>"
        "<html><head><meta charset=\"utf-8\" />"
        f"<style>{base_css}</style>"
        "</head><body>"
        f"{html_content}"
        "</body></html>"
    )


def _parse_indent_to_twips(value: str, default_pt: float = 12.0) -> float | None:
    """
    Parse CSS indent (text-indent / margin-left) to points for Word.
    Returns points (float) or None if unparseable.
    """
    value = (value or "").strip().lower()
    if not value or value == "0" or value == "0px" or value == "0pt":
        return 0.0
    num = re.sub(r"[a-z%]+", "", value)
    if not num:
        return None
    try:
        num_f = float(num)
    except ValueError:
        return None
    if "cm" in value:
        return num_f * 28.35  # 1 cm ≈ 28.35 pt
    if "mm" in value:
        return num_f * 2.835
    if "pt" in value:
        return num_f
    if "px" in value:
        return num_f * 72 / 96  # 96px = 1in = 72pt
    if "em" in value or "rem" in value:
        return num_f * default_pt
    if "in" in value:
        return num_f * 72
    # default assume px
    return num_f * 72 / 96


class _HtmlToDocxWithIndent:
    """
    Wrapper that uses htmldocx but applies text-indent as first_line_indent
    so template paragraph indentation matches the preview in Word/Google Docs.
    """

    def __init__(self):
        from htmldocx import HtmlToDocx
        self._parser = HtmlToDocx()

    def add_html_to_document(self, html: str, document) -> None:
        # Monkey-patch add_styles_to_paragraph on the parser so we can add text-indent
        original = self._parser.add_styles_to_paragraph
        from docx.shared import Pt

        def _add_styles(style: dict) -> None:
            original(style)
            if not self._parser.paragraph:
                return
            if "text-indent" in style:
                pt_val = _parse_indent_to_twips(style["text-indent"])
                if pt_val is not None and pt_val != 0:
                    self._parser.paragraph.paragraph_format.first_line_indent = Pt(pt_val)
        self._parser.add_styles_to_paragraph = _add_styles
        try:
            self._parser.add_html_to_document(html, document)
        finally:
            self._parser.add_styles_to_paragraph = original


def assembled_html_to_docx_bytes(
    html_content: str,
    template_css: str | None = None,
    template_url: str | None = None,
) -> bytes:
    """
    Convert assembled HTML (with SECTION_BREAK markers) to DOCX bytes.
    Format matches static preview: A4, 2.54cm margins, Times New Roman 12pt.
    Same conversion for Google Docs (assembler upload) and DOCX download.
    When template_css is provided, inlines it so DOCX/Google Docs match static preview format.
    When template_url is provided, fetches styles from the template document and merges with
    template_css so alignment and indentation match the uploaded template exactly.
    Strips <style> blocks after inlining; strips citation/source blocks.
    """
    from docx import Document
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
    from docx.shared import Pt, Cm, RGBColor

    from services.assembled_doc_clean import strip_citations_for_assembled

    # Remove citations so download never shows citation lists (filename, Page N.)
    html_content = strip_citations_for_assembled(html_content or "")
    merged_css = _get_merged_template_css(template_css, template_url)
    attempts = []

    html_with_css = html_content
    if merged_css:
        html_with_css = _inline_template_css(html_with_css, merged_css)
    html_with_css = _sanitize_html_for_docx(_ensure_anchor_href(_strip_style_tags(html_with_css)))
    google_ready_html = assembled_html_to_google_import_html(html_content, template_css=template_css, template_url=template_url)
    attempts.append(("altchunk_html", google_ready_html))
    attempts.append(("rich_html", html_with_css))

    html_without_css = _sanitize_html_for_docx(_ensure_anchor_href(_strip_style_tags(html_content)))
    if html_without_css != html_with_css:
        attempts.append(("sanitized_html", html_without_css))

    attempts.append(("plain_text", _plain_text_html(html_without_css or html_content)))

    last_error = None
    for attempt_name, candidate_html in attempts:
        try:
            if attempt_name == "altchunk_html":
                return _build_altchunk_docx(candidate_html)

            document = Document()

            style = document.styles["Normal"]
            style.font.name = FONT_NAME
            style.font.size = Pt(FONT_SIZE_PT)
            style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            style.paragraph_format.space_after = Pt(0)

            for s in document.styles:
                if hasattr(s, "font") and s.font is not None:
                    try:
                        s.font.name = FONT_NAME
                    except Exception:
                        pass

            try:
                section = document.sections[0]
                section.page_height = Cm(PAGE_HEIGHT_CM)
                section.page_width = Cm(PAGE_WIDTH_CM)
                section.left_margin = Cm(MARGIN_CM)
                section.right_margin = Cm(MARGIN_CM)
                section.top_margin = Cm(MARGIN_CM)
                section.bottom_margin = Cm(MARGIN_CM)
            except (IndexError, Exception) as _sec_err:
                logger.warning("assembled_html_to_docx_bytes: could not set page dimensions (%s), proceeding anyway", _sec_err)

            parser = _HtmlToDocxWithIndent()
            content_parts = candidate_html.split(SECTION_BREAK_MARKER)

            for i, part in enumerate(content_parts):
                clean_part = part.strip()
                if not clean_part:
                    continue
                parser.add_html_to_document(clean_part, document)
                if i < len(content_parts) - 1:
                    document.add_page_break()

            black = RGBColor(0, 0, 0)
            for paragraph in document.paragraphs:
                for run in paragraph.runs:
                    run.font.name = FONT_NAME
                    run.font.color.rgb = black
                pf = paragraph.paragraph_format
                pf.right_indent = Pt(0)
                if paragraph.alignment is None:
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

            buf = io.BytesIO()
            document.save(buf)
            buf.seek(0)
            if attempt_name != "rich_html":
                logger.warning("DOCX export succeeded after fallback path: %s", attempt_name)
            return buf.read()
        except Exception as exc:
            last_error = exc
            logger.warning("DOCX export attempt failed (%s): %s", attempt_name, exc)

    if last_error:
        raise last_error
    raise RuntimeError("DOCX export failed with no conversion attempts available")
