"""
Document translation pipeline (PDF/DOCX) built on the Z.AI Translation Agent.

Flow:
  upload bytes → extract (Document AI OCR for PDF, python-docx for DOCX)
  → split into marker-tagged translation chunks → translate chunks concurrently
  via app/services/zai_translation_service.py → reassemble → reconstruct the
  output document → downloadable bytes.

Reconstruction fidelity:
  • DOCX in  → DOCX out: paragraphs (incl. table cells, nested tables) are
    rewritten in place, so styles, headings and table structure survive.
    Inline formatting spans inside a paragraph collapse to the first run's style.
  • PDF in   → PDF out (exact layout): each original page is rasterized as the
    background and translated text is drawn over white patches at the original
    positions (_render_overlay_pdf) — ruling lines, stamps, signatures and
    untranslated regions keep their original pixels. Falls back to a clean
    court-styled restructured PDF when geometry is unavailable.
  • PDF in   → DOCX out: rebuilt from the reconstructed layout (paragraphs +
    tables) as an editable court-styled document via markdown_to_court_docx.

Jobs are tracked in an in-memory store (same pattern as DocumentProcessingQueue
in folder_service.py): restart loses pending jobs, results are held in memory
and expire after JOB_TTL_SECONDS.
"""
from __future__ import annotations

import asyncio
import io
import json
import logging
import re
import tempfile
import threading
import time
import uuid
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from app.services.zai_translation_service import (
    TranslationServiceError,
    translate_text,
)

logger = logging.getLogger("agentic_document_service.document_translation")

MAX_UPLOAD_BYTES = 50 * 1024 * 1024
# Ceiling on translatable characters per document (~150 dense pages) so a single
# job cannot burn unbounded Z.AI tokens.
MAX_TRANSLATABLE_CHARS = 600_000
# One translation call carries at most this many characters / segments. Sized
# so a call finishes in one model pass without truncation, while halving the
# number of round-trips a large document needs; a failed batch still falls back
# to per-segment calls, so accuracy is unaffected by batch size.
CHUNK_CHAR_BUDGET = 3_000
CHUNK_MAX_SEGMENTS = 24


def _concurrency() -> int:
    """Parallel Z.AI calls per job — same settings-driven pool pattern as the
    Document AI page-batch workers in adapters/ocr.py."""
    try:
        from app.core.config import get_settings

        return max(1, int(get_settings().zai_translation_concurrency))
    except Exception:
        return 12
# Per-call timeouts for document mode (payloads are small; a slow call should
# fail fast and be retried, not stall the whole job for minutes).
BATCH_CALL_TIMEOUT_S = 120.0
SEGMENT_CALL_TIMEOUT_S = 75.0
# Overall ceiling on the translating stage: base + per-segment allowance. When
# it trips, the job finishes with whatever translated and reports the rest as
# failed — a partial document beats a stuck job.
STAGE_TIMEOUT_BASE_S = 120.0
STAGE_TIMEOUT_PER_SEGMENT_S = 8.0
STAGE_TIMEOUT_MAX_S = 2_400.0
JOB_TTL_SECONDS = 2 * 60 * 60
MAX_JOBS = 60

_MARKER_RE = re.compile(r"⟦S(\d+)⟧")
# A segment with no letters (page numbers, rules, "____") is passed through untranslated.
_HAS_LETTERS_RE = re.compile(r"[^\W\d_]", re.UNICODE)

_MARKER_GUIDANCE = (
    "The text contains segment markers like ⟦S12⟧ on their own lines. "
    "Keep every marker line exactly as it is, in the same order, and translate only "
    "the text between markers. Do not add, remove, merge or renumber markers."
)


class DocumentTranslationError(Exception):
    def __init__(self, message: str, status_code: int = 400) -> None:
        super().__init__(message)
        self.status_code = status_code


# ── Job store ─────────────────────────────────────────────────────────────────

@dataclass
class TranslationJob:
    job_id: str
    filename: str
    target_lang: str
    source_lang: str
    strategy: str
    output_format: str  # "docx" | "pdf"
    suggestion: str = ""
    status: str = "queued"  # queued|extracting|translating|rebuilding|completed|error
    progress: float = 0.0
    error: str = ""
    segments_total: int = 0
    segments_done: int = 0
    segments_failed: int = 0  # left in the source language after all retries
    page_count: int = 0
    created_at: float = field(default_factory=time.time)
    result_bytes: bytes | None = None
    result_filename: str = ""
    result_mime: str = ""

    def public_state(self) -> dict[str, Any]:
        return {
            "job_id": self.job_id,
            "filename": self.filename,
            "status": self.status,
            "progress": round(min(100.0, max(0.0, self.progress)), 1),
            "error": self.error,
            "segments_total": self.segments_total,
            "segments_done": self.segments_done,
            "segments_failed": self.segments_failed,
            "page_count": self.page_count,
            "target_lang": self.target_lang,
            "source_lang": self.source_lang,
            "strategy": self.strategy,
            "output_format": self.output_format,
            "result_filename": self.result_filename,
            "download_ready": self.result_bytes is not None,
        }


_jobs: dict[str, TranslationJob] = {}
_jobs_lock = threading.Lock()
# Hold task references so the event loop's GC never drops a running pipeline.
_background_tasks: set[asyncio.Task] = set()


def _purge_stale_jobs() -> None:
    now = time.time()
    with _jobs_lock:
        stale = [jid for jid, job in _jobs.items() if now - job.created_at > JOB_TTL_SECONDS]
        for jid in stale:
            _jobs.pop(jid, None)
        if len(_jobs) > MAX_JOBS:
            for jid, _ in sorted(_jobs.items(), key=lambda kv: kv[1].created_at)[: len(_jobs) - MAX_JOBS]:
                _jobs.pop(jid, None)


# Completed results are spooled to disk so a finished translation survives a
# service restart (the in-memory store does not). In-flight jobs are still lost
# on restart — the frontend's 404 handling asks the user to re-upload.
_SPOOL_DIR = Path(tempfile.gettempdir()) / "agentic_translation_results"
_SPOOL_TTL_SECONDS = 24 * 60 * 60
_JOB_ID_RE = re.compile(r"^[0-9a-f]{8,32}$")


def _spool_paths(job_id: str) -> tuple[Path, Path]:
    return _SPOOL_DIR / f"{job_id}.json", _SPOOL_DIR / f"{job_id}.bin"


def _spool_result(job: TranslationJob) -> None:
    """Best-effort: persist a completed job's output; never fail the job over it."""
    try:
        _SPOOL_DIR.mkdir(parents=True, exist_ok=True)
        meta_path, bin_path = _spool_paths(job.job_id)
        bin_path.write_bytes(job.result_bytes or b"")
        meta_path.write_text(
            json.dumps(
                {
                    "filename": job.filename,
                    "target_lang": job.target_lang,
                    "source_lang": job.source_lang,
                    "strategy": job.strategy,
                    "output_format": job.output_format,
                    "result_filename": job.result_filename,
                    "result_mime": job.result_mime,
                    "page_count": job.page_count,
                    "segments_total": job.segments_total,
                    "segments_failed": job.segments_failed,
                }
            )
        )
        now = time.time()
        for stale in _SPOOL_DIR.iterdir():
            try:
                if now - stale.stat().st_mtime > _SPOOL_TTL_SECONDS:
                    stale.unlink()
            except OSError:
                pass
    except Exception as exc:
        logger.warning("[DocTranslation] could not spool result for %s: %s", job.job_id, exc)


def _load_spooled_job(job_id: str) -> TranslationJob | None:
    if not _JOB_ID_RE.match(job_id or ""):
        return None
    meta_path, bin_path = _spool_paths(job_id)
    if not meta_path.is_file() or not bin_path.is_file():
        return None
    try:
        meta = json.loads(meta_path.read_text())
        job = TranslationJob(
            job_id=job_id,
            filename=str(meta.get("filename") or "document"),
            target_lang=str(meta.get("target_lang") or ""),
            source_lang=str(meta.get("source_lang") or "auto"),
            strategy=str(meta.get("strategy") or "general"),
            output_format=str(meta.get("output_format") or "docx"),
        )
        job.status = "completed"
        job.progress = 100.0
        job.result_filename = str(meta.get("result_filename") or "")
        job.result_mime = str(meta.get("result_mime") or "application/octet-stream")
        job.page_count = int(meta.get("page_count") or 0)
        job.segments_total = int(meta.get("segments_total") or 0)
        job.segments_done = job.segments_total
        job.segments_failed = int(meta.get("segments_failed") or 0)
        job.result_bytes = bin_path.read_bytes()
        return job
    except Exception as exc:
        logger.warning("[DocTranslation] could not load spooled job %s: %s", job_id, exc)
        return None


def get_job(job_id: str) -> TranslationJob | None:
    with _jobs_lock:
        job = _jobs.get(job_id)
    if job is not None:
        return job
    job = _load_spooled_job(job_id)
    if job is not None:
        with _jobs_lock:
            _jobs[job_id] = job
    return job


# ── Public API ────────────────────────────────────────────────────────────────

def start_document_translation(
    data: bytes,
    filename: str,
    *,
    target_lang: str,
    source_lang: str = "auto",
    strategy: str = "general",
    output_format: str = "docx",
    suggestion: str = "",
) -> TranslationJob:
    """Validate the upload, register a job, and kick off the async pipeline."""
    _purge_stale_jobs()

    if not data:
        raise DocumentTranslationError("Empty file upload.")
    if len(data) > MAX_UPLOAD_BYTES:
        raise DocumentTranslationError(
            f"File exceeds the {MAX_UPLOAD_BYTES // (1024 * 1024)} MB limit.", status_code=413
        )

    kind = _detect_kind(data, filename)
    output_format = (output_format or "docx").strip().lower()
    if output_format not in {"docx", "pdf"}:
        raise DocumentTranslationError("output_format must be 'docx' or 'pdf'.")
    if kind == "docx" and output_format == "pdf":
        raise DocumentTranslationError(
            "PDF output is only available for PDF uploads; DOCX uploads are translated in place as DOCX."
        )
    if output_format == "pdf":
        from app.services.branding_pdf_service import is_pdf_renderer_available

        if not is_pdf_renderer_available():
            raise DocumentTranslationError(
                "PDF output is unavailable on this server (Chromium renderer not installed); choose DOCX.",
                status_code=503,
            )

    job = TranslationJob(
        job_id=uuid.uuid4().hex[:16],
        filename=filename or f"document.{kind}",
        target_lang=(target_lang or "").strip(),
        source_lang=(source_lang or "auto").strip() or "auto",
        strategy=(strategy or "general").strip().lower(),
        output_format=output_format,
        suggestion=(suggestion or "").strip(),
    )
    if not job.target_lang:
        raise DocumentTranslationError("target_lang is required.")

    with _jobs_lock:
        _jobs[job.job_id] = job

    task = asyncio.get_running_loop().create_task(_run_job(job, data, kind))
    _background_tasks.add(task)
    task.add_done_callback(_background_tasks.discard)
    return job


# ── Pipeline ──────────────────────────────────────────────────────────────────

async def _run_job(job: TranslationJob, data: bytes, kind: str) -> None:
    try:
        job.status = "extracting"
        job.progress = 3.0

        seg_geo: list[dict[str, Any] | None] | None = None
        if kind == "docx":
            docx_doc, segments = await asyncio.to_thread(_extract_docx_segments, data)
            model: Any = docx_doc
            job.progress = 30.0
        else:
            model, segments, seg_geo = await _extract_pdf_segments(job, data)

        translatable = sum(len(s) for i, s in enumerate(segments) if _needs_translation(segments[i]))
        if translatable == 0:
            raise DocumentTranslationError("No translatable text found in the document.", status_code=422)
        if translatable > MAX_TRANSLATABLE_CHARS:
            raise DocumentTranslationError(
                f"Document has ~{translatable:,} characters of text; the limit per job is "
                f"{MAX_TRANSLATABLE_CHARS:,}. Split the document and retry.",
                status_code=413,
            )

        job.status = "translating"
        job.progress = 35.0
        translated = await _translate_segments(job, segments)

        job.status = "rebuilding"
        job.progress = 92.0
        if kind == "docx":
            result = await asyncio.to_thread(_rebuild_docx, model, translated)
            job.result_mime = (
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            ext = "docx"
        else:
            if job.output_format == "pdf":
                # Exact-layout mode: translated text drawn at the original
                # positions over the original page images, so ruling lines,
                # stamps, signatures and letterhead survive untouched.
                result = None
                boxed = sum(1 for g in seg_geo or [] if g)
                if seg_geo and boxed >= max(3, len(seg_geo) // 2):
                    try:
                        result = await asyncio.to_thread(
                            _render_overlay_pdf, data, seg_geo, translated, segments
                        )
                    except Exception as exc:
                        logger.warning(
                            "[DocTranslation] overlay render failed (%s) — falling back to restructured PDF",
                            exc,
                        )
                if result is None:
                    markdown = _pdf_model_to_markdown(model, translated)
                    result = await asyncio.to_thread(_render_pdf, job.filename, markdown)
                job.result_mime = "application/pdf"
                ext = "pdf"
            else:
                markdown = _pdf_model_to_markdown(model, translated)
                result = await asyncio.to_thread(_render_docx_from_markdown, job.filename, markdown)
                job.result_mime = (
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                ext = "docx"

        stem = re.sub(r"\.[A-Za-z0-9]+$", "", job.filename) or "document"
        job.result_filename = f"{stem}_{job.target_lang}.{ext}"
        job.result_bytes = result
        job.status = "completed"
        job.progress = 100.0
        _spool_result(job)
        logger.info(
            "[DocTranslation] job %s done: %s segments, %s -> %s, %d bytes",
            job.job_id, job.segments_total, job.source_lang, job.target_lang, len(result),
        )
    except (DocumentTranslationError, TranslationServiceError) as exc:
        job.status = "error"
        job.error = str(exc)
        logger.warning("[DocTranslation] job %s failed: %s", job.job_id, exc)
    except Exception as exc:  # pragma: no cover
        job.status = "error"
        job.error = "Document translation failed unexpectedly."
        logger.exception("[DocTranslation] job %s crashed: %s", job.job_id, exc)


# ── Extraction ────────────────────────────────────────────────────────────────

def _detect_kind(data: bytes, filename: str) -> str:
    from app.services.adapters.word import looks_like_docx_zip

    name = (filename or "").lower()
    if name.endswith(".pdf") or data[:5] == b"%PDF-":
        return "pdf"
    if name.endswith((".docx", ".doc")):
        if looks_like_docx_zip(data):
            return "docx"
        raise DocumentTranslationError(
            "Legacy binary .doc files are not supported — save as .docx and retry."
        )
    raise DocumentTranslationError("Unsupported file type. Upload a PDF or DOCX document.")


def _iter_docx_paragraphs(container: Any) -> list[Any]:
    """All paragraphs in reading order, descending into tables (and nested tables)."""
    out: list[Any] = list(container.paragraphs)
    for table in container.tables:
        for row in table.rows:
            for cell in row.cells:
                out.extend(_iter_docx_paragraphs(cell))
    return out


def _extract_docx_segments(data: bytes) -> tuple[Any, list[str]]:
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover
        raise DocumentTranslationError("python-docx is not installed on the server.", 500) from exc
    try:
        doc = Document(io.BytesIO(data))
    except Exception as exc:
        raise DocumentTranslationError("Could not open the DOCX file (corrupt or unsupported).") from exc

    paragraphs = [p for p in _iter_docx_paragraphs(doc) if p.text and p.text.strip()]
    # Stash the paragraph list on the document object so rebuild sees the same order.
    doc._translation_paragraphs = paragraphs  # type: ignore[attr-defined]
    return doc, [p.text for p in paragraphs]


async def _extract_pdf_segments(
    job: TranslationJob, data: bytes
) -> tuple[list[dict[str, Any]], list[str], list[dict[str, Any] | None]]:
    """
    OCR the PDF (Document AI; digital-native pages read their embedded text layer)
    and flatten the structured layout into (page model, segment list, geometry).

    The third element is aligned with the segment list: {"page": idx, "bbox": {...}}
    per segment when its position on the original page is known (drives the
    layout-preserving PDF overlay), None otherwise.
    """
    from app.services.adapters.ocr import extract_text_from_bytes

    def _progress(value: float) -> None:
        job.progress = value

    ocr = await asyncio.to_thread(
        extract_text_from_bytes,
        data,
        "application/pdf",
        job.filename,
        progress_callback=_progress,
        progress_start=5.0,
        progress_end=30.0,
        include_tokens=True,
    )
    job.page_count = ocr.page_count

    structured = ocr.structured_json or {}
    pages_json = structured.get("pages") or []
    segments: list[str] = []
    seg_geo: list[dict[str, Any] | None] = []
    model: list[dict[str, Any]] = []

    def _add_segment(
        text: str,
        page_idx: int | None = None,
        bbox: dict[str, float] | None = None,
        tbox: dict[str, float] | None = None,
    ) -> int:
        segments.append(text)
        seg_geo.append(
            {"page": page_idx, "bbox": bbox, "tbox": tbox or bbox}
            if page_idx is not None and bbox
            else None
        )
        return len(segments) - 1

    if pages_json:
        for page_idx, page in enumerate(pages_json):
            blocks: list[dict[str, Any]] = []
            tables = page.get("tables") or []

            # Preferred path: the OCR processor gives line bounding boxes but no
            # tables — rebuild the visual layout (incl. tables) from geometry.
            if not tables:
                layout_blocks = _reconstruct_page_blocks(
                    page.get("lines") or [], page.get("tokens") or []
                )
                if layout_blocks:
                    for block in layout_blocks:
                        if block["kind"] == "para":
                            text = _normalize_ws(block["text"])
                            if text:
                                blocks.append(
                                    {
                                        "type": "para",
                                        "seg": _add_segment(
                                            text, page_idx, block.get("bbox"), block.get("tbox")
                                        ),
                                    }
                                )
                        else:
                            grid = [row for row in block["grid"] if any(c["text"] for c in row)]
                            if not grid:
                                continue
                            header = [
                                [
                                    _add_segment(c["text"], page_idx, c.get("bbox"), c.get("tbox"))
                                    for c in grid[0]
                                ]
                            ]
                            body = [
                                [
                                    _add_segment(c["text"], page_idx, c.get("bbox"), c.get("tbox"))
                                    for c in row
                                ]
                                for row in grid[1:]
                            ]
                            blocks.append({"type": "table", "header": header, "body": body})
                    model.append({"blocks": blocks})
                    continue

            table_text_blob = " ".join(
                _cell_text(cell)
                for table in tables
                for rows in (table.get("headerRows") or [], table.get("bodyRows") or [])
                for row in rows
                for cell in row
            )
            table_blob_norm = _normalize_ws(table_text_blob)

            paragraphs = page.get("paragraphs") or page.get("lines") or []
            for para in paragraphs:
                text = _normalize_ws(str(para.get("text") or ""))
                if not text:
                    continue
                # Document AI paragraphs repeat table content; drop those to avoid duplicates.
                if tables and len(text) > 3 and text in table_blob_norm:
                    continue
                bbox = _line_box(para)
                blocks.append({"type": "para", "seg": _add_segment(text, page_idx, bbox)})

            for table in tables:
                header = [
                    [_add_segment(_cell_text(cell)) for cell in row]
                    for row in (table.get("headerRows") or [])
                ]
                body = [
                    [_add_segment(_cell_text(cell)) for cell in row]
                    for row in (table.get("bodyRows") or [])
                ]
                if header or body:
                    blocks.append({"type": "table", "header": header, "body": body})

            model.append({"blocks": blocks})
    else:
        # No layout at all (pure fallback extraction): paragraph-split the flat text.
        for block in re.split(r"\n\s*\n", ocr.text or ""):
            text = _normalize_ws(block)
            if text:
                model.append({"blocks": [{"type": "para", "seg": _add_segment(text)}]})
        model = [{"blocks": [b for p in model for b in p["blocks"]]}] if model else []

    return model, segments, seg_geo


# ── Geometric layout reconstruction (tables from OCR line boxes) ──────────────
# The configured Document AI OCR processor fills pages[].lines with normalized
# bounding boxes but leaves pages[].tables empty (table extraction needs a
# Form/Layout parser). Scanned court documents are dominated by tables, so we
# rebuild them from geometry: cluster lines into visual rows, detect column
# x-positions, merge wrapped cell continuations, and emit table grids.

_ROW_Y_TOL = 0.008          # rows closer than this (fraction of page height) merge
_CELL_GAP = 0.02            # horizontal gap that separates two cells in a row
_WIDE_ROW_CELL_GAP = 0.014  # stricter merge gap on rows spanning most of the page
_COL_LEFT_TOL = 0.025       # left edges within this cluster into one column
_MIN_TABLE_ROWS = 2         # header + one body row is already a table
_TABLE_ROW_MAX_GAP = 0.045  # vertical gap (fraction of page) that ends a table run


def _line_box(line: dict[str, Any]) -> dict[str, float] | None:
    box = line.get("boundingBox")
    if not isinstance(box, dict):
        return None
    try:
        return {
            "left": float(box["left"]),
            "top": float(box["top"]),
            "right": float(box["left"]) + float(box.get("width") or 0.0),
            "bottom": float(box["top"]) + float(box.get("height") or 0.0),
        }
    except (KeyError, TypeError, ValueError):
        return None


def _cluster_rows(lines: list[dict[str, Any]]) -> list[list[dict[str, Any]]]:
    """Group lines into visual rows by vertical center, top to bottom."""
    boxed = []
    for line in lines:
        box = _line_box(line)
        text = _normalize_ws(str(line.get("text") or ""))
        if box and text:
            boxed.append({"text": text, **box, "cy": (box["top"] + box["bottom"]) / 2})
    boxed.sort(key=lambda l: l["cy"])

    rows: list[list[dict[str, Any]]] = []
    for line in boxed:
        if rows:
            row = rows[-1]
            row_cy = sum(l["cy"] for l in row) / len(row)
            row_height = max(l["bottom"] for l in row) - min(l["top"] for l in row)
            # Cap the tolerance: without it, large display type (letterheads)
            # inflates row_height and chain-merges stacked lines into one "row".
            tol = max(_ROW_Y_TOL, min(row_height * 0.6, 0.014))
            if abs(line["cy"] - row_cy) <= tol:
                row.append(line)
                continue
        rows.append([line])
    for row in rows:
        row.sort(key=lambda l: l["left"])
    return rows


def _row_cells(row: list[dict[str, Any]]) -> list[dict[str, Any]]:
    """Merge horizontally-adjacent lines of a row; keep gap-separated ones as cells."""
    # Letterhead rows that span most of the page carry genuinely separate fields
    # ("Bill No: … * Tal. … Web site: …") — a STRICTER merge gap (merge only when
    # nearly touching) keeps them from being glued into one run-on segment.
    row_span = row[-1]["right"] - row[0]["left"] if row else 0.0
    gap = _WIDE_ROW_CELL_GAP if row_span > 0.8 else _CELL_GAP
    cells: list[dict[str, Any]] = []
    for line in row:
        if cells and line["left"] - cells[-1]["right"] < gap:
            cells[-1]["text"] = f"{cells[-1]['text']} {line['text']}"
            cells[-1]["right"] = max(cells[-1]["right"], line["right"])
            cells[-1]["top"] = min(cells[-1]["top"], line["top"])
            cells[-1]["bottom"] = max(cells[-1]["bottom"], line["bottom"])
        else:
            cells.append(
                {
                    "text": line["text"],
                    "left": line["left"],
                    "right": line["right"],
                    "top": line["top"],
                    "bottom": line["bottom"],
                }
            )
    return cells


def _detect_columns(rows: list[list[dict[str, Any]]]) -> list[float]:
    """Cluster cell left-edges of multi-cell rows into column positions."""
    lefts = sorted(
        cell["left"] for row in rows if len(row) >= 2 for cell in row
    )
    columns: list[list[float]] = []
    for left in lefts:
        if columns and left - columns[-1][-1] <= _COL_LEFT_TOL:
            columns[-1].append(left)
        else:
            columns.append([left])
    return [sum(c) / len(c) for c in columns]


def _assign_to_columns(cells: list[dict[str, Any]], columns: list[float]) -> list[str]:
    out = [""] * len(columns)
    for cell in cells:
        idx = min(range(len(columns)), key=lambda i: abs(columns[i] - cell["left"]))
        out[idx] = f"{out[idx]} {cell['text']}".strip() if out[idx] else cell["text"]
    return out


_GUTTER_STEP = 0.004
_GUTTER_MIN_WIDTH = 0.008
_GUTTER_OPEN_FRAC = 0.85


def _find_column_boundaries(token_rows: list[list[dict[str, Any]]]) -> list[float]:
    """
    Column separators by gap voting: an x-band is a boundary when almost no row
    has a token crossing it. Voting (instead of demanding a fully empty gutter)
    survives the odd long token that leaks across a column gap. Each boundary is
    placed at the middle of the band's widest FULLY-open stretch, so a band that
    runs on through a sparsely-populated column (e.g. an almost-empty "remarks"
    column crossed by 1 of 26 rows) still splits at the true gutter.
    """
    all_tokens = [t for row in token_rows for t in row]
    left = min(t["left"] for t in all_tokens)
    right = max(t["right"] for t in all_tokens)
    n_rows = len(token_rows)

    boundaries: list[float] = []
    band_start: float | None = None
    full_runs: list[tuple[float, float]] = []  # fully-open (frac==1) sub-bands
    full_start: float | None = None

    def _close_band(x_end: float) -> None:
        nonlocal band_start, full_start
        if full_start is not None:
            full_runs.append((full_start, x_end))
            full_start = None
        if band_start is not None and x_end - band_start >= _GUTTER_MIN_WIDTH:
            if full_runs:
                widest = max(full_runs, key=lambda r: r[1] - r[0])
                boundaries.append((widest[0] + widest[1]) / 2)
            else:
                boundaries.append((band_start + x_end) / 2)
        band_start = None
        full_runs.clear()

    x = left + _GUTTER_STEP
    while x < right:
        crossed = sum(
            1
            for row in token_rows
            if any(t["left"] - 0.002 < x < t["right"] + 0.002 for t in row)
        )
        open_frac = 1.0 - crossed / n_rows
        if open_frac >= _GUTTER_OPEN_FRAC:
            if band_start is None:
                band_start = x
            if open_frac >= 1.0 and full_start is None:
                full_start = x
            elif open_frac < 1.0 and full_start is not None:
                full_runs.append((full_start, x))
                full_start = None
        else:
            _close_band(x)
        x += _GUTTER_STEP
    _close_band(x)  # a trailing open band still separates a sparse last column
    return boundaries


def _grid_from_tokens(tokens: list[dict[str, Any]]) -> list[list[dict[str, Any]]] | None:
    """
    Split a table region into cells by projecting word boxes: OCR merges LINES
    across column gaps, but word tokens keep their true x-positions, so the
    whitespace gutters between columns stay empty and define the cell boundaries.
    """
    boxed = []
    for token in tokens:
        box = _line_box(token)
        text = _normalize_ws(str(token.get("text") or ""))
        if box and text:
            boxed.append({"text": text, **box, "cy": (box["top"] + box["bottom"]) / 2})
    if len(boxed) < 4:
        return None

    # Cluster tokens into visual rows (same approach as lines, token-height tolerance).
    boxed.sort(key=lambda t: t["cy"])
    rows: list[list[dict[str, Any]]] = []
    for token in boxed:
        if rows:
            row = rows[-1]
            row_cy = sum(t["cy"] for t in row) / len(row)
            row_height = max(t["bottom"] for t in row) - min(t["top"] for t in row)
            if abs(token["cy"] - row_cy) <= max(_ROW_Y_TOL, min(row_height * 0.6, 0.014)):
                row.append(token)
                continue
        rows.append([token])

    boundaries = _find_column_boundaries(rows)
    n_cols = len(boundaries) + 1
    if n_cols < 2:
        return None

    def _column_of(token: dict[str, Any]) -> int:
        center = (token["left"] + token["right"]) / 2
        for idx, boundary in enumerate(boundaries):
            if center < boundary:
                return idx
        return n_cols - 1

    def _union(a: dict[str, float] | None, b: dict[str, float]) -> dict[str, float]:
        if a is None:
            return dict(b)
        return {
            "left": min(a["left"], b["left"]),
            "top": min(a["top"], b["top"]),
            "right": max(a["right"], b["right"]),
            "bottom": max(a["bottom"], b["bottom"]),
        }

    # Horizontal extent of each column (inset off the gutters so overlay patches
    # never paint over the table's ruling lines). Translated text gets the whole
    # cell width — not just the original words' tight box — so it can render at
    # a readable size instead of shrinking to fit a narrow word cluster.
    region_left = min(t["left"] for t in boxed)
    region_right = max(t["right"] for t in boxed)
    inset = 0.004
    col_spans = []
    for idx in range(n_cols):
        span_l = (boundaries[idx - 1] if idx > 0 else region_left - inset) + inset
        span_r = (boundaries[idx] if idx < n_cols - 1 else region_right + inset) - inset
        col_spans.append((span_l, span_r))

    grid: list[list[dict[str, Any]]] = []
    fold_counts: list[int] = []  # continuation lines folded into each grid row
    for row in rows:
        cells: list[dict[str, Any]] = [
            {"text": "", "bbox": None, "tbox": None} for _ in range(n_cols)
        ]
        for token in sorted(row, key=lambda t: t["left"]):
            col = _column_of(token)
            cell = cells[col]
            cell["text"] = f"{cell['text']} {token['text']}".strip() if cell["text"] else token["text"]
            token_box = {
                "left": token["left"],
                "top": token["top"],
                "right": token["right"],
                "bottom": token["bottom"],
            }
            # tbox = tight box of the actual words (drives alignment inference);
            # bbox = patch box widened to the cell's column span (gives the
            # translation the same room the original cell had).
            cell["tbox"] = _union(cell["tbox"], token_box)
            span_l, span_r = col_spans[col]
            cell["bbox"] = _union(
                cell["bbox"],
                {
                    "left": min(span_l, token["left"]),
                    "top": token["top"],
                    "right": max(span_r, token["right"]),
                    "bottom": token["bottom"],
                },
            )
        # Wrapped continuation: first column empty, sparse, and VERTICALLY
        # ADJACENT to the previous row → fold up. The adjacency test keeps
        # distinct fields further down the page (letterhead numbers, headings)
        # from being swallowed into unrelated cells above them.
        populated = sum(1 for c in cells if c["text"])
        row_top = min(t["top"] for t in row)
        prev_bottom = (
            max(c["tbox"]["bottom"] for c in grid[-1] if c["tbox"]) if grid else 0.0
        )
        if (
            grid
            and not cells[0]["text"]
            and populated <= max(1, n_cols // 2)
            and row_top - prev_bottom < 0.012
            and fold_counts[-1] < 3  # real wraps are 1-3 lines; letterheads chain forever
        ):
            fold_counts[-1] += 1
            for k, cell in enumerate(cells):
                if cell["text"]:
                    prev = grid[-1][k]
                    prev["text"] = f"{prev['text']} {cell['text']}".strip()
                    if cell["bbox"]:
                        prev["bbox"] = _union(prev["bbox"], cell["bbox"])
                    if cell["tbox"]:
                        prev["tbox"] = _union(prev["tbox"], cell["tbox"])
        else:
            grid.append(cells)
            fold_counts.append(0)

    if len(grid) < 2:
        return None
    _merge_wordgap_columns(grid)
    # Drop columns that ended up empty in every row (artifacts of over-splitting).
    n_cols = len(grid[0])
    keep = [k for k in range(n_cols) if any(row[k]["text"] for row in grid)]
    if len(keep) < 2:
        return None
    return [[row[k] for k in keep] for row in grid]


_WORDGAP_MERGE_MAX = 0.012  # tbox gap below this in every co-occupied row = same cell


def _merge_wordgap_columns(grid: list[list[dict[str, Any]]]) -> None:
    """
    Collapse adjacent column pairs that are really ONE column split at a word
    space. Gap voting over-splits sparse columns (e.g. a "Region group" column
    filled in 3 of 15 rows): with almost every row empty there, the space
    between its own words is an open band and becomes a phantom boundary,
    splitting "Middle East" into two cells that then translate word-by-word.
    Real neighbouring columns sit a gutter apart; a split word sits a word
    space apart — merge a pair only when every row occupying both cells shows
    a word-space gap.
    """
    changed = True
    while changed:
        changed = False
        for k in range(len(grid[0]) - 1):
            gaps = [
                row[k + 1]["tbox"]["left"] - row[k]["tbox"]["right"]
                for row in grid
                if row[k]["text"] and row[k + 1]["text"]
            ]
            if not gaps or max(gaps) >= _WORDGAP_MERGE_MAX:
                continue
            for row in grid:
                left_cell, right_cell = row[k], row.pop(k + 1)
                if not right_cell["text"]:
                    continue
                if left_cell["text"]:
                    left_cell["text"] = f"{left_cell['text']} {right_cell['text']}"
                else:
                    left_cell["text"] = right_cell["text"]
                for key in ("bbox", "tbox"):
                    a, b = left_cell[key], right_cell[key]
                    left_cell[key] = (
                        {
                            "left": min(a["left"], b["left"]),
                            "top": min(a["top"], b["top"]),
                            "right": max(a["right"], b["right"]),
                            "bottom": max(a["bottom"], b["bottom"]),
                        }
                        if a and b
                        else (a or b)
                    )
            changed = True
            break


def _tokens_in_band(tokens: list[dict[str, Any]], top: float, bottom: float) -> list[dict[str, Any]]:
    out = []
    for token in tokens:
        box = _line_box(token)
        if box and top - 0.004 <= (box["top"] + box["bottom"]) / 2 <= bottom + 0.004:
            out.append(token)
    return out


def _reconstruct_page_blocks(
    lines: list[dict[str, Any]],
    tokens: list[dict[str, Any]] | None = None,
) -> list[dict[str, Any]] | None:
    """
    Rebuild a page as paragraph/table blocks from line geometry, refining table
    regions with word-token projection when token boxes are available.
    Returns None when the page carries no usable bounding boxes.

    Block shapes: {"kind": "para", "text": str} and
                  {"kind": "table", "grid": list[list[str]]}.
    """
    rows = _cluster_rows(lines)
    if not rows:
        return None

    cell_rows = [_row_cells(row) for row in rows]
    row_bands = [
        (min(l["top"] for l in row), max(l["bottom"] for l in row)) for row in rows
    ]

    # Split the page into runs of table-ish rows vs. plain rows. A single-cell
    # row stays inside an open table run when it is indented past the first
    # column (a wrapped cell continuation). A vertical gap larger than a few
    # row heights ends the run: without this, a distant multi-cell line (page
    # footer, next section's table) joins the run, its word positions corrupt
    # the column detection with phantom columns, and the empty phantom first
    # column then makes the fold heuristic merge real data rows together.
    blocks: list[dict[str, Any]] = []
    i = 0
    while i < len(cell_rows):
        if len(cell_rows[i]) >= 2:
            run = [cell_rows[i]]
            j = i + 1
            first_col_left = cell_rows[i][0]["left"]
            while j < len(cell_rows):
                cells = cell_rows[j]
                if row_bands[j][0] - row_bands[j - 1][1] > _TABLE_ROW_MAX_GAP:
                    break
                if len(cells) >= 2:
                    run.append(cells)
                    first_col_left = min(first_col_left, cells[0]["left"])
                elif cells and cells[0]["left"] > first_col_left + _COL_LEFT_TOL * 2:
                    run.append(cells)  # indented continuation line
                else:
                    break
                j += 1

            multi = sum(1 for r in run if len(r) >= 2)
            if multi >= _MIN_TABLE_ROWS:
                grid: list[list[dict[str, Any]]] | None = None
                if tokens:
                    # Word-box projection: the only way to split cells that OCR
                    # merged into one line across the column gap. When it finds
                    # no consistent column structure, this region is NOT a table
                    # (letterheads, stamp clusters) — never "rescue" it with the
                    # geometry-less line grid, which glues fields into fake
                    # cells that can't even be positioned for the overlay.
                    region_top = min(row_bands[k][0] for k in range(i, j))
                    region_bottom = max(row_bands[k][1] for k in range(i, j))
                    grid = _grid_from_tokens(_tokens_in_band(tokens, region_top, region_bottom))
                else:
                    # No token boxes at all — legacy line-based grid.
                    columns = _detect_columns(run)
                    grid = []
                    for cells in run:
                        assigned = _assign_to_columns(cells, columns)
                        # Wrapped continuation (first column empty, mostly blank
                        # row) folds into the row above.
                        populated = sum(1 for c in assigned if c)
                        if grid and not assigned[0] and populated <= max(1, len(columns) // 2):
                            for k, value in enumerate(assigned):
                                if value:
                                    prev = grid[-1][k]
                                    prev["text"] = f"{prev['text']} {value}".strip()
                        else:
                            grid.append([{"text": value, "bbox": None} for value in assigned])
                if grid is not None and _grid_is_tableish(grid):
                    blocks.append({"kind": "table", "grid": grid})
                else:
                    # Table-like spacing but no consistent column structure —
                    # per-cell paras keep every field positioned and translated.
                    for cells in run:
                        blocks.extend(_cell_para_blocks(cells))
                i = j
                continue
            # Too small to be a table — flush the run one segment PER CELL, never
            # per row: a row-wide segment paints one patch across the whole row,
            # wiping vertical ruling lines and running separate fields together.
            for cells in run:
                blocks.extend(_cell_para_blocks(cells))
            i = j
            continue

        blocks.extend(_cell_para_blocks(cell_rows[i]))
        i += 1

    return blocks


def _grid_is_tableish(grid: list[list[dict[str, Any]]]) -> bool:
    """
    A real table has ≥2 columns, mostly multi-cell rows, AND a mostly-filled
    grid. Letterhead clusters produce sparse many-column grids (fill ~0.5) —
    those must fall back to per-cell paragraphs with tight boxes.
    """
    if not grid or len(grid) < 2 or len(grid[0]) < 2:
        return False
    multi = sum(1 for row in grid if sum(1 for c in row if c["text"]) >= 2)
    if multi < 2 or multi * 2 < len(grid):
        return False
    populated = sum(1 for row in grid for c in row if c["text"])
    return populated > 0.55 * len(grid) * len(grid[0])


def _cell_para_blocks(cells: list[dict[str, Any]]) -> list[dict[str, Any]]:
    out = []
    for cell in cells:
        box = {
            "left": cell["left"],
            "top": cell["top"],
            "right": cell["right"],
            "bottom": cell["bottom"],
        }
        out.append({"kind": "para", "text": cell["text"], "bbox": box, "tbox": dict(box)})
    return out


def _cell_text(cell: Any) -> str:
    if isinstance(cell, dict):
        return _normalize_ws(str(cell.get("text") or ""))
    return _normalize_ws(str(cell or ""))


def _normalize_ws(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip()


# ── Translation ───────────────────────────────────────────────────────────────

def _needs_translation(text: str) -> bool:
    return bool(_HAS_LETTERS_RE.search(text))


# Greek → CJK block: everything decidedly non-Latin (incl. Devanagari & other
# Indic scripts, Arabic, CJK). Latin + Latin-Extended stay outside the range.
_NON_LATIN_RE = re.compile(r"[Ͱ-῿⺀-￿]")
_LATIN_SCRIPT_TARGETS = {"en", "es", "fr", "de", "pt", "it", "nl", "tr", "id", "vi"}


def _looks_untranslated(translated: str, target_lang: str) -> bool:
    """
    Detect provider ECHOES: under load the translation API sometimes returns the
    source text unchanged, which otherwise counts as success and ships original-
    language text in the output. Only decidable when the target uses Latin
    script and the answer still carries non-Latin text.
    """
    if target_lang.split("-")[0].lower() not in _LATIN_SCRIPT_TARGETS:
        return False
    return len(_NON_LATIN_RE.findall(translated or "")) >= 4


def _build_batches(segments: list[str]) -> list[list[int]]:
    """Group translatable segment indices into marker-batch payloads."""
    batches: list[list[int]] = []
    current: list[int] = []
    current_chars = 0
    for idx, text in enumerate(segments):
        if not _needs_translation(text):
            continue
        if current and (
            current_chars + len(text) > CHUNK_CHAR_BUDGET or len(current) >= CHUNK_MAX_SEGMENTS
        ):
            batches.append(current)
            current, current_chars = [], 0
        current.append(idx)
        current_chars += len(text)
    if current:
        batches.append(current)
    return batches


async def _translate_segments(job: TranslationJob, segments: list[str]) -> list[str]:
    translated = list(segments)  # non-translatable segments pass through
    batches = _build_batches(segments)
    job.segments_total = sum(len(b) for b in batches)
    job.segments_done = 0
    if not batches:
        return translated

    suggestion = _MARKER_GUIDANCE if len(batches[0]) > 1 or len(batches) > 1 else ""
    if job.suggestion:
        suggestion = f"{job.suggestion}. {suggestion}".strip()

    # The semaphore guards individual API calls (not whole batch bodies), so
    # per-segment retries from one failed batch run in parallel with everything
    # else instead of serially hogging a slot.
    semaphore = asyncio.Semaphore(_concurrency())
    failed: set[int] = set()

    def _tick(done_increment: int) -> None:
        job.segments_done += done_increment
        if job.segments_total:
            job.progress = 35.0 + 55.0 * min(1.0, job.segments_done / job.segments_total)

    async def _call(payload: str, *, hint: str, timeout_s: float) -> str:
        async with semaphore:
            result = await translate_text(
                payload,
                target_lang=job.target_lang,
                source_lang=job.source_lang,
                strategy=job.strategy,
                suggestion=hint or None,
                timeout_s=timeout_s,
            )
            return result["translated_text"]

    async def _translate_one(i: int, *, count_progress: bool) -> None:
        try:
            text = (await _call(segments[i], hint=job.suggestion, timeout_s=SEGMENT_CALL_TIMEOUT_S)).strip()
            if text and not _looks_untranslated(text, job.target_lang):
                translated[i] = text
                failed.discard(i)
            else:
                failed.add(i)
        except TranslationServiceError as exc:
            if exc.status_code == 400:
                raise  # unsupported language etc. — the job can't succeed
            failed.add(i)
        finally:
            if count_progress:
                _tick(1)

    async def _translate_batch(indices: list[int]) -> None:
        payload = "\n".join(f"⟦S{i}⟧\n{segments[i]}" for i in indices)
        parsed: dict[int, str] | None = None
        try:
            answer = await _call(payload, hint=suggestion, timeout_s=BATCH_CALL_TIMEOUT_S)
            parsed = _parse_marked_output(answer, indices)
        except TranslationServiceError as exc:
            if exc.status_code in {400, 422}:
                raise
            logger.warning("[DocTranslation] batch call failed (%s); falling back per-segment", exc)

        if parsed is None:
            parsed = {}
        # Echoed answers (source text returned unchanged) count as missing —
        # otherwise they'd pass as "translated" and ship original-language
        # text in the output (the page-1-still-in-Marathi failure).
        for i in list(parsed):
            if _looks_untranslated(parsed[i], job.target_lang):
                del parsed[i]

        done = [i for i in indices if parsed.get(i)]
        for i in done:
            translated[i] = parsed[i].strip()
        _tick(len(done))

        missing = [i for i in indices if i not in done]
        if missing:
            await asyncio.gather(*(_translate_one(i, count_progress=True) for i in missing))

    async def _run_stage() -> None:
        await asyncio.gather(*(_translate_batch(batch) for batch in batches))
        # Final retry rounds for segments the provider failed on (timeouts,
        # echoes) — parallel but paced between rounds so a congested API gets
        # room to recover.
        for round_no in range(2):
            if not failed:
                break
            await asyncio.sleep(3.0 * (round_no + 1))
            logger.info(
                "[DocTranslation] retry round %d for %d failed segments", round_no + 1, len(failed)
            )
            await asyncio.gather(
                *(_translate_one(i, count_progress=False) for i in sorted(failed))
            )

    # Hard ceiling: when the provider is having a bad day, finish with a
    # partial translation and an honest failed count instead of hanging.
    deadline = min(
        STAGE_TIMEOUT_MAX_S,
        STAGE_TIMEOUT_BASE_S + STAGE_TIMEOUT_PER_SEGMENT_S * job.segments_total,
    )
    try:
        await asyncio.wait_for(_run_stage(), timeout=deadline)
    except asyncio.TimeoutError:
        logger.warning(
            "[DocTranslation] translation stage hit the %ds deadline; continuing with partial result",
            int(deadline),
        )
        for batch in batches:
            for i in batch:
                if translated[i] == segments[i] or _looks_untranslated(translated[i], job.target_lang):
                    failed.add(i)

    job.segments_failed = len(failed)
    if failed:
        logger.warning(
            "[DocTranslation] %d segments remain untranslated after retries", len(failed)
        )
    return translated


def _parse_marked_output(text: str, expected: list[int]) -> dict[int, str] | None:
    """Split translated output back into segments by ⟦S#⟧ markers; None if unusable."""
    matches = list(_MARKER_RE.finditer(text))
    if not matches:
        return None
    found: dict[int, str] = {}
    for pos, match in enumerate(matches):
        idx = int(match.group(1))
        start = match.end()
        end = matches[pos + 1].start() if pos + 1 < len(matches) else len(text)
        found[idx] = text[start:end].strip()
    missing = [i for i in expected if i not in found or not found[i]]
    # Tolerate a few dropped markers by retrying only if too many are gone.
    if len(missing) > max(1, len(expected) // 5):
        return None
    return {i: found[i] for i in expected if i in found and found[i]}


# ── Reconstruction ────────────────────────────────────────────────────────────

def _rebuild_docx(doc: Any, translated: list[str]) -> bytes:
    paragraphs = getattr(doc, "_translation_paragraphs", [])
    for paragraph, text in zip(paragraphs, translated):
        if paragraph.runs:
            paragraph.runs[0].text = text
            for run in paragraph.runs[1:]:
                run.text = ""
        else:
            paragraph.add_run(text)
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _md_escape_cell(text: str) -> str:
    return text.replace("|", "\\|").replace("\n", " ")


def _pdf_model_to_markdown(model: list[dict[str, Any]], translated: list[str]) -> str:
    out: list[str] = []
    for page_no, page in enumerate(model, start=1):
        if page_no > 1:
            out.append("\n---\n")
        for block in page.get("blocks", []):
            if block["type"] == "para":
                out.append(translated[block["seg"]])
                out.append("")
            else:
                header_rows = block.get("header") or []
                body_rows = block.get("body") or []
                if not header_rows and not body_rows:
                    continue
                first, *rest = header_rows or body_rows[:1]
                width = max(len(first), 1)
                out.append("| " + " | ".join(_md_escape_cell(translated[i]) for i in first) + " |")
                out.append("|" + " --- |" * width)
                remaining = (rest + body_rows) if header_rows else body_rows[1:]
                for row in remaining:
                    out.append("| " + " | ".join(_md_escape_cell(translated[i]) for i in row) + " |")
                out.append("")
    return "\n".join(out).strip()


def _render_docx_from_markdown(filename: str, markdown: str) -> bytes:
    from app.services.docx_export import markdown_to_court_docx

    stem = re.sub(r"\.[A-Za-z0-9]+$", "", filename) or "Translated Document"
    return markdown_to_court_docx(markdown, title=stem)


def _render_overlay_pdf(
    pdf_bytes: bytes,
    seg_geo: list[dict[str, Any] | None],
    translated: list[str],
    originals: list[str] | None = None,
) -> bytes:
    """
    Layout-preserving translation render: each original page is rasterized as the
    background, every TRANSLATED text region is covered with a white patch, and
    the translated text is drawn inside that patch at the same position. Ruling
    lines, stamps, signatures, untranslated regions (numbers, dates) and anything
    OCR didn't read keep their original pixels.
    """
    import base64
    import html as _html

    import fitz  # PyMuPDF

    from playwright.sync_api import sync_playwright

    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    if doc.page_count == 0:
        raise ValueError("PDF has no pages")

    first = doc[0].rect
    page_w_mm = first.width * 25.4 / 72
    page_h_mm = first.height * 25.4 / 72
    # PNG keeps thin ruling lines crisp (JPEG smears them into grey mush); a
    # lower DPI on long documents keeps the print payload manageable.
    dpi = 200 if doc.page_count <= 40 else 130

    numeric_re = re.compile(r"^[\d\s,.\-/%:()=*]+$")

    # Visible segments per page (skipping pass-throughs and junk), so ceilings
    # can be computed against what will actually be painted.
    per_page: dict[int, list[tuple[int, dict[str, float], dict[str, float]]]] = {}
    for i, geo in enumerate(seg_geo):
        if not geo or geo.get("page") is None:
            continue
        text = (translated[i] or "").strip()
        box = geo.get("bbox") or {}
        if not text or "left" not in box:
            continue
        original = (originals[i] if originals else "").strip()
        # Keep the original pixels wherever translation changed nothing (serial
        # numbers, dates) and for single-character junk (stray marks OCR
        # misread). Two-character words ("कर") are real text and must translate.
        if original and (text == original or len(original) < 2):
            continue
        tbox = geo.get("tbox") or box
        per_page.setdefault(geo["page"], []).append((i, box, tbox))

    pages_html: list[str] = []
    for page_index in range(doc.page_count):
        page = doc[page_index]
        pix = page.get_pixmap(dpi=dpi)
        img_b64 = base64.b64encode(pix.tobytes("png")).decode()

        divs: list[str] = []
        entries = per_page.get(page_index, [])
        for i, box, tbox in entries:
            text = (translated[i] or "").strip()
            original = (originals[i] if originals else "").strip()

            # Asymmetric patch padding: Devanagari's shirorekha (top connecting
            # bar) sits ABOVE the OCR box and matras hang below — without this
            # the patch erases letter bodies but leaves black bar remnants.
            line_h = min(tbox["bottom"] - tbox["top"], 0.025)
            w = box["right"] - box["left"]
            pad_x = w * 0.01
            left = max(0.0, box["left"] - pad_x)
            top = max(0.0, box["top"] - line_h * 0.30)
            right = min(1.0, box["right"] + pad_x)
            bottom = min(1.0, box["bottom"] + line_h * 0.15)

            # Growth ceiling: the top of the nearest painted segment below this
            # one (with horizontal overlap), so a wrapping translation grows into
            # its own whitespace but never over the next line or table row.
            ceiling = 0.0
            for _, other_box, _ in entries:
                if other_box is box:
                    continue
                if other_box["top"] >= box["bottom"] - 0.002:
                    overlap = min(right, other_box["right"]) - max(left, other_box["left"])
                    if overlap > 0.2 * (right - left):
                        candidate = other_box["top"] - 0.002
                        if ceiling == 0.0 or candidate < ceiling:
                            ceiling = candidate

            # Alignment: infer from where the words sat inside their cell span;
            # numeric cells (amounts) are right-aligned like the source.
            lm = max(0.0, tbox["left"] - box["left"])
            rm = max(0.0, box["right"] - tbox["right"])
            if original and numeric_re.match(original) and rm < lm:
                align = "right"
            elif lm > 0.02 and rm > 0.02 and abs(lm - rm) < 0.35 * (lm + rm):
                align = "center"
            elif box["left"] > 0.18 and box["right"] < 0.82 and abs((box["left"] + box["right"]) / 2 - 0.5) < 0.06:
                align = "center"
            else:
                align = "left"

            divs.append(
                f'<div class="seg" dir="auto" data-maxbottom="{ceiling * 100:.3f}" '
                f'style="left:{left * 100:.3f}%;top:{top * 100:.3f}%;'
                f'width:{(right - left) * 100:.3f}%;height:{(bottom - top) * 100:.3f}%;'
                f'text-align:{align};">{_html.escape(text)}</div>'
            )
        pages_html.append(
            f'<div class="page"><img src="data:image/png;base64,{img_b64}">{"".join(divs)}</div>'
        )
    doc.close()

    html = f"""<!DOCTYPE html><html><head><meta charset="utf-8"><style>
@page {{ size: {page_w_mm:.2f}mm {page_h_mm:.2f}mm; margin: 0; }}
* {{ margin: 0; padding: 0; box-sizing: border-box; }}
.page {{ position: relative; width: {page_w_mm:.2f}mm; height: {page_h_mm:.2f}mm;
        overflow: hidden; page-break-after: always; }}
.page img {{ position: absolute; inset: 0; width: 100%; height: 100%; }}
.seg {{ position: absolute; background: #fff; overflow: hidden; color: #111;
       line-height: 1.14; padding: 0;
       font-family: 'Noto Sans', 'Noto Sans Devanagari', Arial, sans-serif; }}
</style></head><body>{"".join(pages_html)}
<script>
// Fit text like the original: start from the region's glyph height, and when a
// translation wraps, grow the patch downward only as far as the next painted
// segment below (data-maxbottom) — shrink the font past that, never spill.
document.querySelectorAll('.seg').forEach(function (el) {{
  var baseH = el.offsetHeight;
  var pageH = el.parentElement.offsetHeight;
  var maxBottom = parseFloat(el.dataset.maxbottom || '0');
  var ceiling = maxBottom > 0
    ? Math.max(baseH, maxBottom / 100 * pageH - el.offsetTop)
    : baseH * 1.6;
  var fs = Math.min(baseH * 0.62, 15);
  el.style.fontSize = fs + 'px';
  while (fs > 4.5 && (el.scrollHeight > ceiling || el.scrollWidth > el.clientWidth + 1)) {{
    fs *= 0.94;
    el.style.fontSize = fs + 'px';
  }}
  el.style.height = Math.min(Math.max(el.scrollHeight + 1, baseH), ceiling) + 'px';
}});
</script></body></html>"""

    with sync_playwright() as p:
        browser = p.chromium.launch(args=["--no-sandbox", "--disable-dev-shm-usage"])
        try:
            page = browser.new_page()
            page.set_content(html, wait_until="load")
            pdf = page.pdf(
                prefer_css_page_size=True,
                print_background=True,
                margin={"top": "0", "bottom": "0", "left": "0", "right": "0"},
            )
        finally:
            browser.close()
    return pdf


def _render_pdf(filename: str, markdown: str) -> bytes:
    from app.services.branding_pdf_service import html_to_pdf
    from app.services.merged_pdf_service import build_merged_html

    stem = re.sub(r"\.[A-Za-z0-9]+$", "", filename) or "Translated Document"
    html = build_merged_html(stem, [{"answer": markdown}], include_questions=False)
    return html_to_pdf(html)
