"""
Self-contained Markdown → court-styled DOCX for draft-from-template.

The draft model (gemini-3.1-pro) reproduces the uploaded template as Markdown; a
Markdown chat bubble cannot be filed in court, so this module renders that Markdown
into a downloadable Word document formatted to Indian court conventions (Supreme
Court of India filing style): Times New Roman, A4, 2.54 cm (1") margins, 1.5 line
spacing, justified body, bold headings, preserved legal numbering, tables, and a
page-number footer. `[NOT FOUND IN CASE FILE - provide manually]` markers are bolded
so the reviewing attorney can spot every blank.

Deliberately depends ONLY on python-docx (added to this service's venv) — no HTML
pipeline / htmldocx / premailer, so it stays inside agentic-document-service and does
not couple to any sibling service.
"""
from __future__ import annotations

import io
import logging
import re

logger = logging.getLogger("agentic_document_service.docx_export")

FONT_NAME = "Times New Roman"
FONT_SIZE_PT = 12.0
PAGE_HEIGHT_CM = 29.7   # A4
PAGE_WIDTH_CM = 21.0    # A4
MARGIN_CM = 2.54        # 1 inch (SCI spec)
LINE_SPACING = 1.5
NOT_FOUND_MARKER = "[NOT FOUND IN CASE FILE - provide manually]"

# Red "missing field" placeholders the drafter emits:
#   <span style="color:red;font-weight:bold;">[________ FIELD NAME ________]</span>
# Unwrap the span to its inner text, drop any stray span tags, and colour the
# [________ … ________] blank red + bold so missing fields stand out in the Word file.
_RED_SPAN_RE = re.compile(r"<span[^>]*color:\s*red[^>]*>(.*?)</span>", re.IGNORECASE | re.DOTALL)
_ANY_SPAN_RE = re.compile(r"</?span[^>]*>", re.IGNORECASE)
_RED_BLANK_RE = re.compile(r"\[_{2,}[^\]]*_{2,}\]")
_MARKER_SPLIT_RE = re.compile(rf"({re.escape(NOT_FOUND_MARKER)}|\[_{{2,}}[^\]]*_{{2,}}\])")


def _is_table_row(line: str) -> bool:
    s = line.strip()
    return s.startswith("|") and s.count("|") >= 2


def _is_table_separator(line: str) -> bool:
    s = line.strip()
    if "-" not in s or not s.startswith("|"):
        return False
    return re.fullmatch(r"\|?[\s:\-|]+\|?", s) is not None


def _split_cells(row: str) -> list[str]:
    return [c.strip() for c in row.strip().strip("|").split("|")]


def _emit_marker_runs(paragraph, text: str, *, bold: bool, italic: bool) -> None:
    """Add runs for `text`. NOT_FOUND markers → bold; red [____ FIELD ____] placeholders
    (with or without the <span> wrapper) → bold red so missing fields stand out."""
    from docx.shared import RGBColor
    # Unwrap red placeholder spans to their inner [____ FIELD ____] text; drop any stray tags.
    text = _RED_SPAN_RE.sub(lambda m: m.group(1), text)
    text = _ANY_SPAN_RE.sub("", text)
    if not text:
        return
    # Split keeping NOT_FOUND markers and red blanks as their own tokens.
    for fragment in re.split(r"(<br\s*/?>)", text, flags=re.IGNORECASE):
        if not fragment:
            continue
        if re.fullmatch(r"<br\s*/?>", fragment, flags=re.IGNORECASE):
            paragraph.add_run().add_break()
            continue
        for part in _MARKER_SPLIT_RE.split(fragment):
            if not part:
                continue
            if part == NOT_FOUND_MARKER:
                run = paragraph.add_run(part)
                run.bold = True
            elif _RED_BLANK_RE.fullmatch(part):
                run = paragraph.add_run(part)
                run.bold = True
                run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)  # red — a field to fill
            else:
                run = paragraph.add_run(part)
                run.bold = bold
                run.italic = italic


def _emit_inline(paragraph, text: str, *, base_bold: bool = False) -> None:
    """Render inline Markdown: ** ** toggles bold; NOT_FOUND markers are bolded."""
    # Toggle bold on ** boundaries (unbalanced ** degrade gracefully).
    bold = base_bold
    for segment in text.split("**"):
        if segment:
            _emit_marker_runs(paragraph, segment, bold=bold, italic=False)
        bold = not bold


def _add_heading(doc, text: str, level: int, *, center: bool = False) -> None:
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10 if level <= 1 else 8)
    pf.space_after = Pt(4)
    pf.keep_with_next = True
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _emit_inline(p, text, base_bold=True)
    size = 14.0 if level <= 1 else (13.0 if level == 2 else 12.0)
    for run in p.runs:
        run.bold = True
        run.font.size = Pt(size)


# Lines that are ALWAYS centered in Indian legal documents.
_CENTER_ALWAYS_RE = re.compile(r"^(VERSUS|VS\.?|V/S\.?)$", re.IGNORECASE)
# Trailing party-role labels ("...Plaintiff", "…LANDLORD/LESSOR") are right-aligned.
_ROLE_LABEL_RE = re.compile(
    r"^[.…\s]*[.…]\s*[\"'“”]?(the\s+)?(first|second|third|1st|2nd|3rd)?\s*"
    r"(plaintiff|defendant|petitioner|respondent|appellant|applicant|complainant|"
    r"landlord|tenant|lessor|lessee|licensor|licensee|vendor|purchaser|party|witness(es)?)"
    r"(\s*/\s*\w+)?[\"'“”]?\s*\.?$",
    re.IGNORECASE,
)
# A Markdown horizontal rule (--- / *** / ___, 3+ chars) — used to separate the
# executable agreement from the trailing "ITEMS REQUIRING COMPLETION" / notes sections.
_MD_HR_RE = re.compile(r"^\s*([-*_])(\s*\1){2,}\s*$")


def _add_hr(doc) -> None:
    """Insert a thin full-width horizontal rule (a bottom-bordered empty paragraph)."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "808080")
    borders.append(bottom)
    p_pr.append(borders)


def _looks_like_layout_table(rows: list[list[str]]) -> bool:
    if not rows or max((len(r) for r in rows), default=0) != 2:
        return False
    header = " ".join(rows[0]).lower()
    body = " ".join(" ".join(r) for r in rows[:4]).lower()
    if "witness 1" in header and "witness 2" in header:
        return True
    if "signature" in body and any(a in body for a in ("landlord", "lessor", "licensor", "vendor")) and any(
        b in body for b in ("tenant", "lessee", "licensee", "purchaser")
    ):
        return True
    return False


def _remove_table_borders(table) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tbl_pr = table._tbl.tblPr
    for child in list(tbl_pr):
        if child.tag == qn("w:tblBorders"):
            tbl_pr.remove(child)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        borders.append(el)
    tbl_pr.append(borders)


def _add_table(doc, rows: list[list[str]], *, layout: bool = False) -> None:
    if not rows:
        return
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.shared import Pt

    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    if layout:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        _remove_table_borders(table)
    else:
        try:
            table.style = "Table Grid"  # built-in bordered style
        except Exception:
            pass
    for r_idx, row_cells in enumerate(rows):
        for c_idx in range(n_cols):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            para = cell.paragraphs[0]
            para.text = ""
            para.paragraph_format.space_after = Pt(2 if layout else 6)
            _emit_inline(para, row_cells[c_idx] if c_idx < len(row_cells) else "",
                         base_bold=(r_idx == 0))
    doc.add_paragraph()  # spacer after table


def _add_page_number_footer(doc) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    footer = doc.sections[0].footer
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    run._r.append(begin); run._r.append(instr); run._r.append(end)


def _add_draft_header(doc) -> None:
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    header = doc.sections[0].header
    para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run("DRAFT — for attorney review; verify every entry before filing")
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


def markdown_to_court_docx(
    markdown_text: str, *, title: str | None = None, typography: dict | None = None
) -> bytes:
    """Render draft Markdown into a court-styled DOCX and return its bytes.

    typography (optional): the drafting pipeline's Stage-A structural analysis, e.g.
    ``{"base_font": {"font": "Arial", "size_pt": 11}, "title_format": {...}}``. When
    given, the document's base font/size follow the template instead of the court
    defaults. None (the single-call path) keeps the Times New Roman 12 court style.
    """
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

    md = (markdown_text or "").replace("\r\n", "\n").replace("\r", "\n")
    # A court draft must read as a finished document — strip any residual inline
    # source tags (e.g. " [source: file.pdf]") the model may have emitted.
    md = re.sub(r"[ \t]*\[source:[^\]]*\]", "", md, flags=re.IGNORECASE)
    # Safety net: some models glue the operative-words preamble onto clause 1
    # ("...WITNESSETH AS FOLLOWS:1. TERM:"), which reads as one run-on paragraph in the
    # DOCX. Break after such a colon when a numbered clause immediately follows so the
    # preamble and clause 1 land in separate paragraphs.
    md = re.sub(
        r"(WITNESSETH\b[^\n:]{0,80}:|\bFOLLOWS\s*:)\s*(?=\d{1,2}\.)",
        r"\1\n\n",
        md,
        flags=re.IGNORECASE,
    )

    doc = Document()

    # Page: A4 with 1-inch (2.54 cm) margins.
    section = doc.sections[0]
    section.page_height = Cm(PAGE_HEIGHT_CM)
    section.page_width = Cm(PAGE_WIDTH_CM)
    section.left_margin = Cm(MARGIN_CM)
    section.right_margin = Cm(MARGIN_CM)
    section.top_margin = Cm(MARGIN_CM)
    section.bottom_margin = Cm(MARGIN_CM)

    # Normal style: Times New Roman 12, 1.5 spacing, justified — unless the template's
    # own base font/size was captured by the drafting pipeline (typography.base_font).
    _font_name, _font_size = FONT_NAME, FONT_SIZE_PT
    _base = (typography or {}).get("base_font") if isinstance(typography, dict) else None
    if isinstance(_base, dict):
        if str(_base.get("font") or "").strip():
            _font_name = str(_base["font"]).strip()
        try:
            _sz = float(_base.get("size_pt") or 0)
            if 8.0 <= _sz <= 20.0:  # ignore nonsense values, keep it readable/filable
                _font_size = _sz
        except (TypeError, ValueError):
            pass
    normal = doc.styles["Normal"]
    normal.font.name = _font_name
    normal.font.size = Pt(_font_size)
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = LINE_SPACING
    pf.space_after = Pt(6)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    if title:
        try:
            doc.core_properties.title = str(title)[:200]
        except Exception:
            pass

    _add_draft_header(doc)
    _add_page_number_footer(doc)

    try:
        from docx.enum.text import WD_ALIGN_PARAGRAPH as _ALIGN
        lines = md.split("\n")
        i = 0
        # Indian legal-document alignment: the opening title block (document title,
        # stamp-paper note, court name / suit number / IN THE MATTER OF lines) is
        # CENTERED until the first body paragraph; "VERSUS" is always centered;
        # trailing party-role labels ("...Plaintiff", "…LANDLORD") are right-aligned.
        in_title_block = True
        while i < len(lines):
            line = lines[i]
            stripped = line.strip()
            if not stripped:
                i += 1
                continue

            # Horizontal rule (--- / *** / ___) → thin full-width separator. Used to
            # divide the executable agreement from the trailing notes; also ends the
            # centered title block.
            if _MD_HR_RE.match(stripped):
                _add_hr(doc)
                in_title_block = False
                i += 1
                continue

            # Table block: a row followed by a separator row.
            if _is_table_row(line) and i + 1 < len(lines) and _is_table_separator(lines[i + 1]):
                block_rows = [_split_cells(line)]
                j = i + 2  # skip the separator
                while j < len(lines) and _is_table_row(lines[j]):
                    block_rows.append(_split_cells(lines[j]))
                    j += 1
                _add_table(doc, block_rows, layout=_looks_like_layout_table(block_rows))
                in_title_block = False
                i = j
                continue

            plain = stripped.replace("**", "").strip()

            # ATX heading (#, ##, ###).
            if stripped.startswith("#"):
                hashes = len(stripped) - len(stripped.lstrip("#"))
                text = stripped.lstrip("#").strip()
                center = in_title_block or bool(_CENTER_ALWAYS_RE.match(text.replace("**", "").strip()))
                _add_heading(doc, text, hashes, center=center)
                i += 1
                continue

            # Whole-line bold => heading (the model emits **BOLD HEADINGS**).
            if (stripped.startswith("**") and stripped.endswith("**")
                    and stripped.count("**") == 2 and len(stripped) > 4):
                center = in_title_block or bool(_CENTER_ALWAYS_RE.match(plain))
                _add_heading(doc, plain, 2, center=center)
                i += 1
                continue

            # Bare VERSUS / VS. line (not bold) — still centered + bold.
            if _CENTER_ALWAYS_RE.match(plain):
                _add_heading(doc, plain, 2, center=True)
                i += 1
                continue

            # Trailing party-role label => right-aligned bold line.
            if _ROLE_LABEL_RE.match(plain):
                para = doc.add_paragraph()
                para.alignment = _ALIGN.RIGHT
                _emit_inline(para, stripped, base_bold=True)
                i += 1
                continue

            # Parenthetical note directly under the title (stamp-paper line) stays centered.
            if in_title_block and plain.startswith("(") and plain.endswith(")"):
                para = doc.add_paragraph()
                para.alignment = _ALIGN.CENTER
                _emit_inline(para, stripped)
                i += 1
                continue

            # Normal paragraph (preserve legal numbering / bullets as literal text).
            in_title_block = False
            para = doc.add_paragraph()
            _emit_inline(para, stripped)
            i += 1
    except Exception as exc:
        # Never fail the request over formatting — fall back to a plain-text dump.
        logger.warning("markdown_to_court_docx: rich parse failed (%s); using plain fallback", exc)
        doc = Document()
        for raw_line in md.split("\n"):
            doc.add_paragraph(raw_line)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
