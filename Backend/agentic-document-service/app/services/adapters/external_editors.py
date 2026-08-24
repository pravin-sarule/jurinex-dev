"""Native Google Docs / Zoho Writer integrations for Case Storage.

Ported from Backend/drafting-service (Google, per-user OAuth) and
Backend/zoho-service (Zoho Office Integrator) so Case Storage has no runtime
dependency on those Node services.
"""
from __future__ import annotations

import io
import logging
import os
import re
from pathlib import Path

import httpx

# The app's pydantic settings read .env into a settings object, not os.environ —
# this module uses os.environ directly, so load the service .env here (no overrides).
try:
    from dotenv import load_dotenv

    load_dotenv(Path(__file__).resolve().parents[3] / ".env")
except Exception:  # pragma: no cover — dotenv ships with pydantic-settings
    pass

logger = logging.getLogger("agentic_document_service.external_editors")

ZOHO_WRITER_ENDPOINT = "https://api.office-integrator.in/writer/officeapi/v1/document"
GOOGLE_TOKEN_ENDPOINT = "https://oauth2.googleapis.com/token"
GOOGLE_DRIVE_FILES_ENDPOINT = "https://www.googleapis.com/drive/v3/files"

DOCX_MIMETYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
GOOGLE_DOC_MIMETYPE = "application/vnd.google-apps.document"


# ── Zoho Office Integrator ───────────────────────────────────────────────────

def build_blank_docx() -> bytes:
    """Minimal empty DOCX for a new blank Writer document."""
    from docx import Document as DocxDocument

    buf = io.BytesIO()
    DocxDocument().save(buf)
    return buf.getvalue()


def _extract_zoho_document_token(document_url: str | None) -> str | None:
    if not document_url:
        return None
    for pattern in (r"/open/([^/?]+)", r"/documents/([^/?]+)"):
        match = re.search(pattern, document_url)
        if match:
            return match.group(1)
    return None


def create_zoho_writer_session(
    *,
    signed_url: str,
    file_name: str,
    file_id: str,
    save_callback_url: str,
) -> dict:
    """Create a Zoho Writer editing session for a DOCX reachable at signed_url.

    Mirrors zoho-service officeIntegratorService.createWriterSession: multipart
    POST with apikey + permissions + callback settings; Zoho later POSTs the
    saved DOCX to save_callback_url.
    """
    api_key = os.environ.get("ZOHO_OI_API_KEY")
    if not api_key:
        raise RuntimeError("ZOHO_OI_API_KEY is not configured")

    import json as _json

    data = {
        "apikey": api_key,
        "permissions": _json.dumps({
            "document.export": True,
            "document.print": True,
            "document.edit": True,
            "review.changes.resolve": False,
            "review.comment": True,
            "collab.chat": True,
        }),
        "editor_settings": _json.dumps({"unit": "in", "language": "en", "view": "pageview"}),
        "callback_settings": _json.dumps({
            "save_format": "docx",
            "save_url": save_callback_url,
            "context_info": "Case Storage save callback",
        }),
        "document_info": _json.dumps({"document_name": file_name, "document_id": str(file_id)}),
        "url": signed_url,
    }

    # httpx sends `data` as multipart when combined with a files= argument; Zoho's
    # API accepts standard form encoding as well, but mirror the Node service and
    # send multipart for safety.
    files = {key: (None, value) for key, value in data.items()}
    resp = httpx.post(ZOHO_WRITER_ENDPOINT, files=files, timeout=30)
    if resp.status_code != 200:
        raise RuntimeError(f"Zoho Writer error ({resp.status_code}): {resp.text[:300]}")
    payload = resp.json()
    document_url = payload.get("document_url")
    return {
        "type": "zoho",
        "editor": "writer",
        "iframeUrl": document_url,
        "sessionId": payload.get("session_id"),
        "zohoDocumentToken": _extract_zoho_document_token(document_url),
        "documentId": payload.get("document_id"),
    }


def public_base_url() -> str:
    """Base URL Zoho's servers use to reach our save-callback (must be public)."""
    return (
        os.environ.get("PUBLIC_BASE_URL")
        or os.environ.get("SERVICE_PUBLIC_URL")
        or "http://localhost:8092"
    ).rstrip("/")


# ── Google Docs (per-user OAuth, tokens held by auth-service) ────────────────

def _auth_service_url() -> str:
    return (os.environ.get("AUTH_SERVICE_URL") or "http://localhost:5001").rstrip("/")


def fetch_google_refresh_token(user_id: str) -> str | None:
    """Fetch the user's google_drive_refresh_token from authservice.

    Same endpoint drafting-service used (utils/oauth2Client.js:58-68):
    GET /api/auth/internal/user/{id}/tokens → {google_drive_refresh_token, google_drive_token_expiry, email}
    """
    headers = {"X-Internal-Request": "true"}
    internal_token = os.environ.get("INTERNAL_SERVICE_TOKEN")
    if internal_token:
        # Only attach when configured — an empty "Bearer " header trips authservice's JWT parsing.
        headers["Authorization"] = f"Bearer {internal_token}"
    try:
        resp = httpx.get(
            f"{_auth_service_url()}/api/auth/internal/user/{user_id}/tokens",
            headers=headers,
            timeout=15,
        )
        if resp.status_code != 200:
            logger.info("[GoogleDocs] auth-service token lookup %s for user %s", resp.status_code, user_id)
            return None
        body = resp.json() or {}
        return body.get("google_drive_refresh_token")
    except Exception as exc:
        logger.warning("[GoogleDocs] auth-service unreachable: %s", exc)
        return None


def google_drive_status(user_id: str, authorization: str | None) -> dict:
    """Connection status + connect URL.

    Connected = a refresh token exists. authservice's own /status flag also checks
    the stored ACCESS-token expiry, which is routinely in the past — we mint fresh
    access tokens ourselves, so only the refresh token matters here.
    """
    if fetch_google_refresh_token(user_id):
        return {"connected": True}
    status: dict = {"connected": False}
    headers = {"Authorization": authorization} if authorization else {}
    try:
        resp = httpx.get(f"{_auth_service_url()}/api/auth/google/drive", headers=headers, timeout=15)
        body = resp.json() if resp.headers.get("content-type", "").startswith("application/json") else {}
        if resp.status_code == 200:
            status["authUrl"] = (body or {}).get("authUrl")
        else:
            status["error"] = (body or {}).get("message") or f"auth service returned {resp.status_code}"
            logger.warning("[GoogleDocs] authUrl fetch failed (%s): %s", resp.status_code, status["error"])
    except Exception as exc:
        status["error"] = f"auth service unreachable: {exc}"
    return status


def _mint_google_access_token(refresh_token: str) -> str:
    client_id = os.environ.get("GOOGLE_CLIENT_ID")
    client_secret = os.environ.get("GOOGLE_CLIENT_SECRET")
    if not client_id or not client_secret:
        raise RuntimeError("GOOGLE_CLIENT_ID / GOOGLE_CLIENT_SECRET are not configured")
    resp = httpx.post(
        GOOGLE_TOKEN_ENDPOINT,
        data={
            "client_id": client_id,
            "client_secret": client_secret,
            "refresh_token": refresh_token,
            "grant_type": "refresh_token",
        },
        timeout=20,
    )
    if resp.status_code != 200:
        raise RuntimeError(f"Google token refresh failed ({resp.status_code}): {resp.text[:200]}")
    token = (resp.json() or {}).get("access_token")
    if not token:
        raise RuntimeError("Google token refresh returned no access_token")
    return token


def export_google_doc_pdf(user_id: str, google_file_id: str) -> bytes:
    """Export a Google Doc as PDF bytes (fresh snapshot for chat grounding)."""
    refresh_token = fetch_google_refresh_token(user_id)
    if not refresh_token:
        raise PermissionError("GOOGLE_NOT_CONNECTED")
    access_token = _mint_google_access_token(refresh_token)
    resp = httpx.get(
        f"{GOOGLE_DRIVE_FILES_ENDPOINT}/{google_file_id}/export",
        params={"mimeType": "application/pdf"},
        headers={"Authorization": f"Bearer {access_token}"},
        timeout=60,
    )
    if resp.status_code != 200:
        raise RuntimeError(f"Google Doc export failed ({resp.status_code}): {resp.text[:200]}")
    return resp.content


def extract_docx_text(docx_bytes: bytes) -> str:
    """Plain-text extraction from a DOCX (paragraphs + tables) for chat grounding."""
    from docx import Document as DocxDocument

    doc = DocxDocument(io.BytesIO(docx_bytes))
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def create_google_doc(user_id: str, title: str) -> dict:
    """Create a blank Google Doc in the user's Drive. Returns {google_file_id, web_view_link}."""
    refresh_token = fetch_google_refresh_token(user_id)
    if not refresh_token:
        raise PermissionError("GOOGLE_NOT_CONNECTED")
    access_token = _mint_google_access_token(refresh_token)
    resp = httpx.post(
        GOOGLE_DRIVE_FILES_ENDPOINT,
        params={"fields": "id,name,webViewLink"},
        json={"name": title, "mimeType": GOOGLE_DOC_MIMETYPE},
        headers={"Authorization": f"Bearer {access_token}"},
        timeout=30,
    )
    if resp.status_code not in (200, 201):
        raise RuntimeError(f"Google Drive create failed ({resp.status_code}): {resp.text[:200]}")
    body = resp.json() or {}
    return {"google_file_id": body.get("id"), "web_view_link": body.get("webViewLink")}


def google_connected(user_id: str) -> bool:
    return bool(fetch_google_refresh_token(user_id))
