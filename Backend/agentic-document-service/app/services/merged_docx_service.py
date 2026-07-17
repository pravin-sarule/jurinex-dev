"""
Build a production-quality .docx from selected Q&A sections.

The answers are GitHub-Flavored Markdown (the platform's LLM output contract),
so this module renders a pragmatic GFM subset natively into python-docx
elements: headings (#–######), paragraphs, bullet/numbered lists, pipe tables,
horizontal rules, and inline **bold** / *italic* / `code`.
"""
from __future__ import annotations

import io
import re
from datetime import datetime, timezone
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

_BOX_CHARS_RE = re.compile(r"[┌└├┤┬┴┼│─┐┘]")
# Whole decorative banner boxes (e.g. "⚖️ LEXIS ..." headers) are dropped
# entirely — their inner text is branding noise, not answer content.
_BOX_BLOCK_RE = re.compile(r"┌[\s─]*┐[\s\S]*?└[\s─]*┘")
_HTML_TAG_RE = re.compile(r"</?[a-zA-Z][^>]*>")
_INLINE_TOKEN_RE = re.compile(
    r"(\*\*[^*\n]+\*\*"      # **bold**
    r"|\*[^*\n]+\*"          # *italic*
    r"|`[^`\n]+`)"           # `code`
)
_TABLE_ROW_RE = re.compile(r"^\s*\|.*\|\s*$")
_TABLE_SEPARATOR_RE = re.compile(r"^\s*\|(\s*:?-{2,}:?\s*\|)+\s*$")
_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
_BULLET_RE = re.compile(r"^\s*[-*+]\s+(.*)$")
_NUMBERED_RE = re.compile(r"^\s*\d+[.)]\s+(.*)$")
_HR_RE = re.compile(r"^\s*([-*_]\s*){3,}$")


def _clean_answer_markdown(text: str) -> str:
    """Strip UI-only artifacts (ASCII banners, raw HTML tags) from an answer."""
    text = _BOX_BLOCK_RE.sub("", str(text or ""))
    lines = []
    for line in str(text or "").splitlines():
        stripped = _BOX_CHARS_RE.sub(" ", line) if _BOX_CHARS_RE.search(line) else line
        if _BOX_CHARS_RE.search(line) and not stripped.strip():
            continue
        lines.append(stripped)
    cleaned = "\n".join(lines)
    cleaned = _HTML_TAG_RE.sub("", cleaned)
    cleaned = cleaned.replace("&nbsp;", " ").replace("&amp;", "&").replace("&lt;", "<").replace("&gt;", ">")
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
    return cleaned.strip()


def _add_inline_runs(paragraph, text: str) -> None:
    """Render **bold** / *italic* / `code` spans as styled runs."""
    for token in _INLINE_TOKEN_RE.split(text):
        if not token:
            continue
        if token.startswith("**") and token.endswith("**") and len(token) > 4:
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("*") and token.endswith("*") and len(token) > 2:
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        elif token.startswith("`") and token.endswith("`") and len(token) > 2:
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        else:
            paragraph.add_run(token)


def _split_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def _add_markdown_table(document, rows: list[str]) -> None:
    header = _split_table_row(rows[0])
    body_lines = [r for r in rows[1:] if not _TABLE_SEPARATOR_RE.match(r)]
    n_cols = max(len(header), *(len(_split_table_row(r)) for r in body_lines)) if body_lines else len(header)
    table = document.add_table(rows=1 + len(body_lines), cols=n_cols)
    table.style = "Table Grid"
    for col, cell_text in enumerate(header):
        cell = table.rows[0].cells[col]
        cell.text = ""
        _add_inline_runs(cell.paragraphs[0], cell_text)
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for row_index, line in enumerate(body_lines, start=1):
        for col, cell_text in enumerate(_split_table_row(line)[:n_cols]):
            cell = table.rows[row_index].cells[col]
            cell.text = ""
            _add_inline_runs(cell.paragraphs[0], cell_text)


def _render_markdown(document, markdown: str, *, heading_offset: int = 1) -> None:
    """
    Render a GFM subset into the document.

    heading_offset shifts markdown heading levels so answer headings nest under
    the numbered section heading (## in the answer -> Heading 3 in the doc).
    """
    lines = markdown.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]

        if not line.strip():
            i += 1
            continue

        # Fenced code block
        if line.lstrip().startswith("```"):
            code_lines: list[str] = []
            i += 1
            while i < len(lines) and not lines[i].lstrip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing fence
            para = document.add_paragraph()
            run = para.add_run("\n".join(code_lines))
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            continue

        # Table block
        if _TABLE_ROW_RE.match(line) and i + 1 < len(lines) and _TABLE_SEPARATOR_RE.match(lines[i + 1]):
            table_rows = []
            while i < len(lines) and _TABLE_ROW_RE.match(lines[i]):
                table_rows.append(lines[i])
                i += 1
            _add_markdown_table(document, table_rows)
            document.add_paragraph()
            continue

        heading = _HEADING_RE.match(line)
        if heading:
            level = min(9, len(heading.group(1)) + heading_offset)
            text = heading.group(2).strip().strip("*").strip()
            document.add_heading(text, level=level)
            i += 1
            continue

        if _HR_RE.match(line):
            i += 1
            continue

        bullet = _BULLET_RE.match(line)
        if bullet:
            para = document.add_paragraph(style="List Bullet")
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            _add_inline_runs(para, bullet.group(1))
            i += 1
            continue

        numbered = _NUMBERED_RE.match(line)
        if numbered:
            para = document.add_paragraph(style="List Number")
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            _add_inline_runs(para, numbered.group(1))
            i += 1
            continue

        # Plain paragraph — join soft-wrapped lines until a structural break.
        para_lines = [line.strip()]
        i += 1
        while (
            i < len(lines)
            and lines[i].strip()
            and not _HEADING_RE.match(lines[i])
            and not _BULLET_RE.match(lines[i])
            and not _NUMBERED_RE.match(lines[i])
            and not _TABLE_ROW_RE.match(lines[i])
            and not lines[i].lstrip().startswith("```")
        ):
            para_lines.append(lines[i].strip())
            i += 1
        para = document.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _add_inline_runs(para, " ".join(para_lines))


def _append_field(run, instruction: str) -> None:
    """Append a Word field (e.g. PAGE / NUMPAGES) to a run."""
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def _add_page_number_footer(document) -> None:
    """Centred 'Page X of Y' footer on every page."""
    footer_para = document.sections[0].footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for text, field in (("Page ", "PAGE"), (" of ", "NUMPAGES")):
        run = footer_para.add_run(text)
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
        field_run = footer_para.add_run()
        field_run.font.size = Pt(8.5)
        field_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
        _append_field(field_run, field)


def build_merged_docx(
    title: str,
    sections: list[dict[str, Any]],
    *,
    include_questions: bool = True,
) -> bytes:
    """
    Assemble the selected Q&A sections into a single .docx.

    Each section dict: {question, answer, source?, origin_label?}.
    include_questions=False omits the numbered question headings so the
    document reads as a continuous set of answers.
    """
    document = Document()
    _add_page_number_footer(document)

    doc_title = (title or "").strip() or "Merged Legal Analysis"
    heading = document.add_heading(doc_title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run(
        f"Generated on {datetime.now(timezone.utc).strftime('%d %B %Y')}"
        + (f" · {len(sections)} section(s)" if len(sections) > 1 else "")
    )
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    for index, section in enumerate(sections, start=1):
        if include_questions:
            question = str(section.get("question") or "").strip() or f"Question {index}"
            document.add_heading(f"{index}. {question}", level=1)
        elif index > 1:
            document.add_paragraph()  # breathing room between unlabelled answers

        origin_label = str(section.get("origin_label") or "").strip()
        if origin_label:
            origin_para = document.add_paragraph()
            origin_run = origin_para.add_run(origin_label)
            origin_run.font.size = Pt(8)
            origin_run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
            origin_run.italic = True

        answer = _clean_answer_markdown(str(section.get("answer") or ""))
        if answer:
            _render_markdown(document, answer, heading_offset=1)
        else:
            document.add_paragraph("(No answer content)")

        source = str(section.get("source") or "").strip()
        if source:
            source_para = document.add_paragraph()
            source_run = source_para.add_run(f"Source: {source}")
            source_run.font.size = Pt(8)
            source_run.italic = True
            source_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
